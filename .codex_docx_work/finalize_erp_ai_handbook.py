from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path(r"D:\PythonProject\LearnOne\ERP与AI融合实践-讲解知识手册-完善版.docx")
OUTPUT_PATH = Path(r"D:\PythonProject\LearnOne\ERP与AI融合实践-讲解知识手册-会前终版.docx")


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"未找到段落：{exact_text}")


def find_table(document, text_fragment):
    for table in document.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if text_fragment in table_text:
            return table
    raise ValueError(f"未找到表格：{text_fragment}")


def add_paragraph_after(document, anchor_element, text="", style=None):
    paragraph_element = OxmlElement("w:p")
    anchor_element.addnext(paragraph_element)
    paragraph = Paragraph(paragraph_element, document._body)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def add_labeled_paragraph_after(document, anchor_element, label, body):
    paragraph = add_paragraph_after(document, anchor_element, style="Normal")
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(body)
    return paragraph


def clone_table_after(document, anchor_element, template_table):
    table_element = deepcopy(template_table._tbl)
    anchor_element.addnext(table_element)
    return Table(table_element, document._body)


def set_callout(table, label, body):
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(f"  {body}")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def keep_heading(paragraph):
    paragraph.paragraph_format.keep_with_next = True


document = Document(SOURCE_PATH)
callout_template = find_table(document, "核心结论")


# 一、加入可直接照读的5分钟精简讲稿。
anchor = find_table(document, "如果时间有限")._tbl
heading = add_paragraph_after(document, anchor, "5分钟精简讲稿（可直接照读）", "Heading 2")
keep_heading(heading)
script_parts = [
    (
        "开场：",
        "今天分享的重点不是模型有多聪明，而是我们如何让AI真正进入ERP业务流程。传统ERP依赖菜单和表单，用户需要理解系统结构；我们希望用户先表达业务目标，再由AI帮助完成信息收集和交互。",
    ),
    (
        "审批场景：",
        "用户可以直接说想请假或发起其他审批。系统先识别意图，再读取当前公司可用的审批模板，根据模板动态追问缺失字段。字段完整后生成预览，只有用户明确确认，Workflow才调用ERP Tool写入。",
    ),
    (
        "日报场景：",
        "用户不需要逐项填写日报表单，只要描述今天完成的工作、遇到的问题和明天计划，AI就能整理成结构化草稿。用户可以继续修改、取消或确认，正式提交仍然受流程和权限控制。",
    ),
    (
        "技术选择：",
        "LangChain提供模型、Prompt、Tool和结构化输出等基础能力；LangGraph负责有状态、多步骤、可中断和可恢复的业务流程；Deep Agents用于探索更开放的自主规划。ERP写入风险较高，因此主流程优先选择更可控的LangGraph。",
    ),
    (
        "安全边界：",
        "AI并不直接操作数据库，也不能绕过ERP权限。模型负责理解和建议，Schema负责约束数据结构，Workflow负责校验、预览和确认，ERP负责最终规则和可信写入，LangSmith等观测能力帮助我们定位问题。",
    ),
    (
        "当前结论：",
        "项目已经验证审批和日报从自然语言输入到ERP写入的完整闭环，但这还不等于已经具备全面上线条件。下一步需要继续补齐权限验证、幂等、异常恢复、评测集、指标基线和灰度运营。",
    ),
    (
        "收束：",
        "所以，AI与ERP融合的关键不是给AI无限自主权，而是把自然语言理解与企业确定性流程结合起来，在安全、可控、可追踪的前提下，让用户更自然地使用ERP。",
    ),
]
anchor_element = heading._p
for label, body in script_parts:
    paragraph = add_labeled_paragraph_after(document, anchor_element, label, body)
    anchor_element = paragraph._p
usage_callout = clone_table_after(document, anchor_element, callout_template)
set_callout(
    usage_callout,
    "使用建议",
    "先把这段完整读两遍，再回到正文补技术细节。现场忘词时，只要按“场景—技术—安全—边界—结论”五步继续即可。",
)


# 二、加入两个场景的完整演示口播脚本。
anchor = find_table(document, "演示原则")._tbl
heading = add_paragraph_after(document, anchor, "4. 两个场景的演示口播脚本", "Heading 2")
keep_heading(heading)
approval_heading = add_paragraph_after(document, heading._p, "审批演示口播", "Heading 3")
keep_heading(approval_heading)
approval_steps = [
    "我先用一句自然语言发起审批。示例输入应选择现场环境已经验证过的审批类型，例如：“我明天下午想请半天假，家里有事。”",
    "系统首先判断这是一条审批请求，然后读取当前用户、公司和可用模板；这里不是让模型凭记忆猜模板。",
    "如果模板要求的信息还不完整，系统会根据真实字段定义继续追问。缺什么问什么，而不是一次展示整张表单。",
    "字段完整后先生成预览。我会强调：现在还没有写入ERP，用户仍然可以修改或取消。",
    "用户明确确认后，Workflow才调用受控ERP Tool提交，并返回申请结果。最后一句总结是：入口变简单了，但原有规则没有被绕过。",
]
anchor_element = approval_heading._p
for step in approval_steps:
    paragraph = add_paragraph_after(document, anchor_element, step, "List Number")
    anchor_element = paragraph._p

report_heading = add_paragraph_after(document, anchor_element, "日报演示口播", "Heading 3")
keep_heading(report_heading)
report_steps = [
    "我先描述一段当天工作，例如完成了哪些功能、解决了什么问题以及明天准备做什么，不按照日报字段逐项填写。",
    "AI将自然语言整理为ERP要求的日报结构。这里展示的价值不是“写一段漂亮文字”，而是把非结构化表达转换为可提交的数据。",
    "系统保存草稿并生成预览。如果内容不准确，我可以直接说需要修改哪一部分，流程会保留已有上下文。",
    "确认前流程可以中断，稍后再从Checkpoint恢复；这说明系统管理的不只是一轮回答，而是持续的业务状态。",
    "用户确认后再提交ERP。最后总结：Agent负责理解和组织内容，Workflow负责草稿、预览、确认和写入闸门。",
]
anchor_element = report_heading._p
for step in report_steps:
    paragraph = add_paragraph_after(document, anchor_element, step, "List Number")
    anchor_element = paragraph._p
demo_callout = clone_table_after(document, anchor_element, callout_template)
set_callout(
    demo_callout,
    "每段演示的固定收口",
    "用户少填了哪些步骤、系统保留了哪些规则、后台记录了哪些过程。用这三点收口，比继续解释代码和节点更容易让听众记住。",
)


# 三、补充五层架构口径，帮助回答整体架构问题。
anchor = find_table(document, "LangChain提供基础组件")._tbl
heading = add_paragraph_after(document, anchor, "用五层架构解释整体方案", "Heading 2")
keep_heading(heading)
architecture_items = [
    "第一层，交互入口：用户通过聊天表达请假、日报等业务目标，不再先寻找菜单和理解表单。",
    "第二层，AI理解：模型识别意图、抽取信息、组织内容，并根据缺失字段生成自然语言追问。",
    "第三层，流程编排：LangGraph管理State、节点、分支、Interrupt/Resume和人工确认，保证流程能暂停、恢复和审计。",
    "第四层，ERP能力：ERP通过受控Tool提供模板、字段、权限、校验、草稿和提交能力，继续承担可信执行。",
    "第五层，治理观测：通过Trace、日志、评测和业务指标监控模型判断、流程路由、Tool调用、成本和最终效果。",
]
anchor_element = heading._p
for item in architecture_items:
    paragraph = add_paragraph_after(document, anchor_element, item, "List Bullet")
    anchor_element = paragraph._p
architecture_callout = clone_table_after(document, anchor_element, callout_template)
set_callout(
    architecture_callout,
    "一句话架构",
    "聊天负责入口，模型负责理解，LangGraph负责流程，ERP负责执行，观测体系负责治理。",
)


# 四、对行业数据增加会前核对提醒，不新增未经验证的数据。
anchor = find_paragraph(document, "来源二：https://www.pwc.com/gx/en/issues/c-suite-insights/ceo-survey.html")._p
data_callout = clone_table_after(document, anchor, callout_template)
set_callout(
    data_callout,
    "上台前核对",
    "确认PPT中的报告年份、样本数量、统计口径和脚注与原始来源一致。若无法再次核实具体数字，现场以“AI应用普及，但规模化价值仍需流程和运营支撑”的定性结论为主。",
)


# 五、补充回答追问的固定结构。
anchor = find_table(document, "避免直接说")._tbl
heading = add_paragraph_after(document, anchor, "回答领导追问的四步法", "Heading 2")
keep_heading(heading)
answer_steps = [
    "先给结论：第一句话直接回答“能不能、为什么、目前到哪一步”，不要先铺技术背景。",
    "再讲机制：用一到两个关键机制说明为什么，例如预览确认、Schema校验、权限上下文或Checkpoint。",
    "主动讲边界：明确当前验证范围、尚未完成的生产化能力，以及哪些结论还需要真实数据证明。",
    "最后给下一步：说明准备如何验证、补齐或量化，让回答从解释问题落到行动计划。",
]
anchor_element = heading._p
for step in answer_steps:
    paragraph = add_paragraph_after(document, anchor_element, step, "List Number")
    anchor_element = paragraph._p
answer_callout = clone_table_after(document, anchor_element, callout_template)
set_callout(
    answer_callout,
    "示例",
    "“目前可以进入试点，但还不能直接全面上线；因为核心闭环已跑通，生产所需的权限穿透、幂等和灰度指标仍需补齐；下一步先用真实用户和模板做小范围验证。”",
)


# 六、增加一分钟收尾和上台前检查清单。
anchor = find_table(document, "收尾句")._tbl
heading = add_paragraph_after(document, anchor, "十四、上台前最后检查", "Heading 1")
keep_heading(heading)
closing_heading = add_paragraph_after(document, heading._p, "1. 一分钟收尾话术", "Heading 2")
keep_heading(closing_heading)
closing_callout = clone_table_after(document, closing_heading._p, callout_template)
set_callout(
    closing_callout,
    "可直接照读",
    "这次实践最重要的收获，不是做出了一个会聊天的功能，而是验证了AI可以在受控条件下进入真实ERP流程。用户通过自然语言表达目标，AI降低交互门槛，LangGraph保证流程可控，ERP继续承担规则、权限和可信写入。当前我们已经跑通审批和日报闭环，下一步会围绕统一Schema、统一Tool、评测指标和生产治理继续完善，让试点能力逐步沉淀为可复用的企业AI底座。",
)

check_heading = add_paragraph_after(document, closing_callout._tbl, "2. 会前检查清单", "Heading 2")
keep_heading(check_heading)
check_items = [
    "把30秒开场和一分钟收尾各完整说两遍，确保开头、结尾不依赖临场发挥。",
    "用计时器完整讲一遍，控制在目标时长以内；超时优先删术语定义和行业数据解释。",
    "确认演示账号、公司、权限、审批模板和日报日期均为可用状态。",
    "准备审批和日报的成功截图或短录屏，现场异常时能够立即切换。",
    "清理演示环境中的敏感数据、无关通知、历史聊天和可能暴露个人信息的内容。",
    "再次核对PPT中的行业数据、来源年份、样本数量和脚注，避免把调查比例解释成效率提升比例。",
    "准备回答三类问题：为什么选择这个方案、怎样保证安全、下一步如何证明价值。",
    "讲解时每页只保留一个结论；如果听众已经理解，就向后推进，不为了讲完手册而拖慢节奏。",
]
anchor_element = check_heading._p
for item in check_items:
    paragraph = add_paragraph_after(document, anchor_element, item, "List Bullet")
    anchor_element = paragraph._p
final_callout = clone_table_after(document, anchor_element, callout_template)
set_callout(
    final_callout,
    "最后提醒",
    "不要背完整文档。记住主线和边界，用演示证明闭环，用问答补充技术细节。",
)


for paragraph in document.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        keep_heading(paragraph)

document.core_properties.title = "ERP与AI融合实践：智能审批与日报讲解知识手册（会前终版）"
document.save(OUTPUT_PATH)
print(OUTPUT_PATH)
