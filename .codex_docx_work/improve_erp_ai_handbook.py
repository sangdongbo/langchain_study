from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path(r"D:\PythonProject\LearnOne\ERP与AI融合实践-讲解知识手册.docx")
OUTPUT_PATH = Path(r"D:\PythonProject\LearnOne\ERP与AI融合实践-讲解知识手册-完善版.docx")


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"未找到段落：{exact_text}")


def find_table(document, text_fragment):
    for table in document.tables:
        if text_fragment in "\n".join(cell.text for row in table.rows for cell in row.cells):
            return table
    raise ValueError(f"未找到表格：{text_fragment}")


def add_paragraph_after(document, anchor_element, text, style=None):
    paragraph_element = OxmlElement("w:p")
    anchor_element.addnext(paragraph_element)
    paragraph = Paragraph(paragraph_element, document._body)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def clone_table_after(document, anchor_element, template_table):
    table_element = deepcopy(template_table._tbl)
    anchor_element.addnext(table_element)
    return Table(table_element, document._body)


def set_callout(table, label, body):
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(f"  {body}")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_matrix(table, rows):
    if len(table.rows) != len(rows) or len(table.columns) != len(rows[0]):
        raise ValueError("模板表格尺寸与数据不一致")
    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.bold = row_index == 0
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def mark_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


document = Document(SOURCE_PATH)
callout_template = find_table(document, "核心结论")
three_column_template = find_table(document, "回答重点")
four_column_template = find_table(document, "主要风险")


# 使用方式：增加明天的优先阅读路线。
anchor = find_table(document, "口径提醒")._tbl
heading = add_paragraph_after(document, anchor, "明天优先阅读路线", "Heading 2")
mark_keep_with_next(heading)
route_callout = clone_table_after(document, heading._p, callout_template)
set_callout(
    route_callout,
    "如果时间有限",
    "先读第一、二、四、七、十二和十三节；重点记住核心结论、两个业务闭环、框架选择理由、安全边界、高频问答和八句话总结。",
)


# 主线：增加可直接使用的开场话术。
anchor = find_paragraph(
    document,
    "这次分享不应被讲成“介绍几个 Agent”，而应讲成一次 ERP 业务入口和执行方式的探索：用户用自然语言表达目标，AI根据上下文补充信息，Workflow控制流程，ERP完成校验与写入。",
)._p
opening_callout = clone_table_after(document, anchor, callout_template)
set_callout(
    opening_callout,
    "30秒开场",
    "今天不重点讲模型有多聪明，而是分享我们怎样把AI接入审批和日报两个真实ERP流程。用户用自然语言表达目标，AI负责理解和补充信息，Workflow负责控制步骤，ERP继续负责权限、校验和最终写入。",
)


# 主线：增加15-20分钟时间分配。
anchor = find_table(document, "现场过渡")._tbl
heading = add_paragraph_after(document, anchor, "15-20分钟讲解节奏", "Heading 2")
mark_keep_with_next(heading)
timing_table = clone_table_after(document, heading._p, four_column_template)
set_matrix(
    timing_table,
    [
        ["阶段", "建议时间", "主要内容", "必须说清的一句话"],
        ["开场与业务价值", "2-3分钟", "为什么选择审批、日报；AI改变了什么", "AI增加自然语言入口，但不替代ERP规则。"],
        ["方案与演示", "6-8分钟", "审批闭环、日报闭环、预览确认", "开放理解交给AI，确定性写入交给Workflow。"],
        ["技术选择与难点", "4-5分钟", "LangGraph、状态、动态字段、安全护栏", "选择LangGraph是为了可控、可恢复和可审计。"],
        ["效果、边界与答疑", "3-4分钟", "已验证效果、上线条件、量化指标", "当前验证了闭环，规模化价值需要用真实基线数据证明。"],
    ],
)
timing_note = clone_table_after(document, timing_table._tbl, callout_template)
set_callout(
    timing_note,
    "时间不够时",
    "优先保留业务闭环、架构选择和安全边界；术语定义、行业数据和项目名称速查可以放到答疑阶段。",
)


# 业务场景：补充现场演示步骤与故障兜底。
anchor = find_paragraph(document, "正式提交仍由 Workflow 和 ERP Tool 控制。")._p
heading = add_paragraph_after(document, anchor, "3. 现场演示顺序与备用方案", "Heading 2")
mark_keep_with_next(heading)
demo_items = [
    "先演示审批：输入一句自然语言请求，展示动态追问、预览和确认，不要一开始就讲节点细节。",
    "再演示日报：重点展示自然语言到结构化草稿、修改以及中断恢复，突出两个场景共用的闭环能力。",
    "每个演示结束只总结三点：用户少填了什么、系统守住了什么、后台记录了什么。",
    "若现场接口或模型异常，立即切换到预先准备的截图或录屏，按“输入—追问—预览—确认—ERP结果”复述，不在台上临时排障。",
]
for item in demo_items:
    paragraph = add_paragraph_after(document, heading._p, item, "List Bullet")
    heading = paragraph
demo_callout = clone_table_after(document, heading._p, callout_template)
set_callout(
    demo_callout,
    "演示原则",
    "演示的目标是证明业务闭环，不是证明模型每次都能自由发挥；成功标准是过程可理解、结果可核对、写入有边界。",
)


# 效果章节：明确试点与生产上线的边界。
anchor = find_paragraph(document, "通过节点路径和Trace观察流程执行情况。")._p
heading = add_paragraph_after(document, anchor, "从试点验证到生产上线的边界", "Heading 2")
mark_keep_with_next(heading)
boundary_table = clone_table_after(document, heading._p, three_column_template)
set_matrix(
    boundary_table,
    [
        ["维度", "当前可说明", "上线前必须补齐"],
        ["业务闭环", "已完成追问、草稿/预览、确认和ERP写入路径验证", "覆盖更多真实模板、异常输入和边界场景"],
        ["安全控制", "已采用权限上下文、结构化输出和提交确认思路", "完成权限穿透测试、敏感数据策略和审计留痕"],
        ["稳定性", "已验证中断恢复及Tool调用链路", "补齐超时、重试、幂等、降级、人工接管和告警"],
        ["价值证明", "已证明技术可行和业务流程可走通", "建立传统流程基线，持续跟踪效率、质量、成本和用户采用率"],
    ],
)
boundary_callout = clone_table_after(document, boundary_table._tbl, callout_template)
set_callout(
    boundary_callout,
    "推荐口径",
    "当前结论是“具备试点价值并验证了完整闭环”，不是“已经达到全面生产上线标准”。",
)


# 高频问答：补充数据、稳定性、成本和评估四类追问。
anchor_paragraph = find_paragraph(
    document,
    "可以扩展订单、库存、客户跟进、费用报销、采购申请和经营分析，但应优先选择高频、规则清晰、风险可控且效果可衡量的场景。",
)
anchor_element = anchor_paragraph._p
additional_qas = [
    (
        "业务数据发给模型，会不会产生泄露风险？",
        "风险不能只靠模型厂商承诺解决。正式上线前应明确模型部署和数据处理方式，落实最小数据、字段脱敏、传输与存储控制、权限隔离和审计策略；不需要的敏感字段不进入模型上下文。",
    ),
    (
        "模型判断不稳定，系统如何兜底？",
        "模型只负责理解、抽取和建议，关键结果必须通过结构化输出、Schema校验和业务规则检查。失败时可以重试、回退到固定流程或转人工；未通过预览确认和提交守卫时不能写入ERP。",
    ),
    (
        "如何控制响应时间和运行成本？",
        "应按任务复杂度选择模型，压缩无关上下文，缓存稳定配置，限制无效循环和Tool调用次数，并持续观察单次任务耗时、Token消耗和ERP接口延迟。成本要与节省的人工时间和错误成本一起评估。",
    ),
    (
        "如何证明模型效果可以持续，而不是只在演示中可用？",
        "建立覆盖常见表达、模糊输入、缺失字段和异常场景的评测集，版本变更后重复回归；上线试点后再结合模板识别准确率、人工修正率、完成率、失败原因和用户反馈持续评估。",
    ),
]
for question, answer in additional_qas:
    question_paragraph = add_paragraph_after(document, anchor_element, question, "Heading 2")
    mark_keep_with_next(question_paragraph)
    answer_paragraph = add_paragraph_after(document, question_paragraph._p, answer, "Normal")
    anchor_element = answer_paragraph._p


# 高频问答末尾：增加现场用词边界。
heading = add_paragraph_after(document, anchor_element, "现场表达注意事项", "Heading 2")
mark_keep_with_next(heading)
wording_table = clone_table_after(document, heading._p, three_column_template)
set_matrix(
    wording_table,
    [
        ["避免直接说", "建议表达", "原因"],
        ["AI替代ERP", "AI为ERP增加自然语言入口", "规则、权限和可信写入仍由ERP承担"],
        ["Agent可以自主完成所有业务", "开放判断由Agent辅助，关键动作由Workflow约束", "避免夸大自主性和忽略责任边界"],
        ["系统已经可以全面上线", "当前已验证试点闭环，正在补齐生产化能力", "区分技术验证和生产标准"],
        ["效率提升了固定百分比", "先建立传统流程基线，再用真实指标对比", "没有样本和口径时不应承诺数字"],
    ],
)


# 结构性检查：新内容应使用真实标题、列表和表格，避免悬空标题。
for paragraph in document.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        mark_keep_with_next(paragraph)

document.core_properties.title = "ERP与AI融合实践：智能审批与日报讲解知识手册（完善版）"
document.save(OUTPUT_PATH)
print(OUTPUT_PATH)
