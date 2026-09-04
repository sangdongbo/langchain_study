from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"D:\PythonProject\LearnOne\ai_erp_rag_assistant\docs\ERP_APPROVAL_DYNAMIC_FORM_FRONTEND.docx")


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D9D9", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        p_pr.remove(node)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color="666666")


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Body Text")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_code(doc, value, language="json"):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.0
    # A small label keeps code samples distinguishable without adding a box.
    if language:
        label = p.add_run(f"{language}\n")
        set_run_font(label, name="Consolas", size=8, bold=True, color="35618F")
    run = p.add_run(value)
    set_run_font(run, name="Consolas", size=8.2, color="1F2937")
    p._p.get_or_add_pPr().append(OxmlElement("w:shd"))
    p._p.pPr[-1].set(qn("w:fill"), "F5F7FA")
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, "35618F")
        set_cell_border(cell)
        set_cell_padding(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(value))
        set_run_font(r, size=9.2, bold=True, color="FFFFFF")
        if widths:
            cell.width = Cm(widths[i])
    for ridx, row_values in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_shading(cell, "F3F6FA" if ridx % 2 else "FFFFFF")
            set_cell_border(cell)
            set_cell_padding(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_run_font(r, size=9.1, color="222222")
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size={1: 15, 2: 12.5, 3: 11}[level], bold=True)
    set_keep_with_next(p)
    return p


def json_text(value):
    return json.dumps(value, ensure_ascii=False, indent=2)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)
styles["Body Text"].font.name = "Microsoft YaHei"
styles["Body Text"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Body Text"].font.size = Pt(10.5)
styles["Body Text"].paragraph_format.space_after = Pt(6)
styles["Body Text"].paragraph_format.line_spacing = 1.18
for style_name, size in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)):
    styles[style_name].font.name = "Microsoft YaHei"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles[style_name].font.size = Pt(size)
    styles[style_name].font.bold = True
    styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
styles["Title"].font.name = "Microsoft YaHei"
styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Title"].font.size = Pt(23)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(0, 0, 0)
styles["Title"].paragraph_format.space_after = Pt(4)

code_style = styles.add_style("Code Block", 1)
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
code_style.font.size = Pt(8.2)
code_style.paragraph_format.space_after = Pt(7)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("ERP 审批动态表单前端对接说明  ·  ")
set_run_font(r, size=9, color="666666")
add_page_number(fp)

# Title page
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ERP 审批动态表单前端对接说明")
set_run_font(r, size=23, bold=True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("动态字段渲染 多轮收集 审批人选择 预览确认")
set_run_font(sr, size=11.5, color="35618F")
doc.add_paragraph()
add_body(doc, "这份说明给 ERP 前端页面使用，目标是让页面不依赖某一种固定审批类型，而是根据服务端返回的 form_schema.fields 自动渲染输入框、下拉、单选、多选、人员选择、明细表和附件等控件。文档中的 JSON 按当前 ai_erp_rag_assistant 实际接口和响应模型整理。")
add_body(doc, "最重要的结论：POST /api/chat 返回的是一个扁平的 ChatResponse。回复文本在 message，动态表单在 form_schema，审批预览在 preview；当前没有统一包裹在 content、answer 或 data 里的二次结构。")
add_table(doc, ["前端读取位置", "用途"], [
    ["response.message", "聊天文字和状态提示"],
    ["response.workflow_status", "决定填写、选人、预览还是提交完成"],
    ["response.form_schema.fields", "动态表单字段定义"],
    ["response.form_schema.values", "服务端已收集的当前值"],
    ["response.form_schema.missing_field_keys", "缺少的字段 name"],
    ["response.form_schema.invalid_fields", "字段级错误"],
    ["response.preview", "提交前预览、审批流和幂等信息"],
    ["response.erp_data", "ERP 查询或提交结果"],
], widths=[5.3, 10.8])

add_heading(doc, "1 对接接口和身份", 1)
add_body(doc, "审批页面通常需要以下接口。身份建议放在 HTTP 请求头中，前端不要让普通用户编辑 UID、Authorization 或 company_id。")
add_table(doc, ["接口", "用途", "前端调用时机"], [
    ["POST /api/assistants/list", "获取审批助手和 RAG 助手", "页面初始化"],
    ["POST /api/approval/templates", "查询当前用户可用审批模板", "用户选择或搜索审批类型"],
    ["POST /api/approval/form-schema", "获取指定模板的完整动态字段", "确定模板后或刷新字段"],
    ["POST /api/approval/options", "分页获取人员、部门、订单等选项", "远程选择器搜索时"],
    ["POST /api/chat", "多轮收集、校验、预览和确认提交", "用户发送消息或点击操作按钮"],
    ["POST /api/sessions/list", "读取历史会话列表", "左侧历史记录"],
    ["POST /api/sessions/messages", "读取历史消息和卡片", "切换会话"],
], widths=[5.2, 6.4, 4.5])
add_body(doc, "推荐请求头：")
add_code(doc, "UID: 863\nAuthorization: Bearer <ERP_TOKEN>", "http")
add_body(doc, "审批助手请求体的核心字段如下。assistant_key 固定使用 approval-assistant；同一个审批过程必须复用同一个 session_id。")
add_code(doc, json_text({
    "message": "帮我发起一个请假审批",
    "session_id": "approval-20260904-001",
    "request_id": "req-20260904-001",
    "assistant_key": "approval-assistant",
    "user_id": "863",
    "company_id": "16",
    "form_values": {},
    "selected_assignees": {},
    "confirm": None,
    "preview_id": "",
    "preview_version": None,
    "preview_hash": "",
    "reset": False,
    "stream": False,
}))

add_heading(doc, "2 ChatResponse 完整外层结构", 1)
add_body(doc, "以下是前端应按照的完整外层结构。字段可能为空数组、空对象或 null，但不要改用其他字段名。")
chat_shape = {
    "message": "字段已补齐，已生成审批预览。请回复“确认提交”或“取消”。",
    "route": "approval_workflow",
    "assistant_type": "approval",
    "plan": {},
    "tool_calls": [],
    "errors": [],
    "pending_question": "",
    "workflow_status": "preview_ready",
    "erp_mode": "remote",
    "erp_write_mode": "disabled",
    "evidence": [],
    "citations": [],
    "erp_data": {},
    "form_schema": {},
    "preview": {},
}
add_code(doc, json_text(chat_shape))
add_body(doc, "当前 response 顶层没有 session_id 和 assistant_key。前端应从本次请求上下文中保留这两个值；历史消息接口返回的 response 也不会替代前端自己的会话标识。")

add_heading(doc, "3 form_schema 字段协议", 1)
add_body(doc, "form_schema 是动态表单的唯一渲染依据。前端不要解析 ERP 原始 field_type，也不要通过 label 猜控件；优先使用 component、value_type、ui、validation 和 option_source。")
add_code(doc, json_text({
    "schema_version": "1.0",
    "template": {
        "template_id": "101",
        "template_code": "101",
        "title": "综合审批示例",
        "company_id": "16",
    },
    "fields": [
        {
            "name": "reason",
            "label": "申请事由",
            "required": True,
            "type": "text",
            "component": "text",
            "value_type": "string",
            "erp_field_type": "input",
            "input_type": "",
            "options": [],
            "option_values": [],
            "option_source": None,
            "validation": {"max_length": 200},
            "group": None,
            "group_key": None,
            "group_label": None,
            "group_type": None,
            "ui": {
                "placeholder": "请输入申请事由",
                "multiple": False,
                "readonly": False,
                "hidden": False,
                "col_span": None,
            },
            "children": [],
            "erp": {
                "field_key": "reason",
                "field_id": "",
                "field_type": "input",
                "sort": 1,
            },
        },
        {
            "name": "expense_type",
            "label": "费用类型",
            "required": True,
            "type": "enum",
            "component": "select",
            "value_type": "string",
            "erp_field_type": "select",
            "input_type": "",
            "options": ["差旅费", "办公费", "招待费"],
            "option_values": [
                {"label": "差旅费", "value": "travel", "disabled": False, "meta": {}},
                {"label": "办公费", "value": "office", "disabled": False, "meta": {}},
                {"label": "招待费", "value": "entertainment", "disabled": False, "meta": {}},
            ],
            "option_source": {"type": "static", "lazy": False, "searchable": False},
            "validation": {},
            "ui": {
                "placeholder": "请选择费用类型",
                "multiple": False,
                "readonly": False,
                "hidden": False,
                "col_span": None,
            },
            "children": [],
        },
        {
            "name": "notify_users",
            "label": "抄送人员",
            "required": False,
            "type": "array",
            "component": "entity-select",
            "value_type": "array",
            "erp_field_type": "user",
            "options": [],
            "option_values": [],
            "option_source": {"type": "user_list", "lazy": True, "searchable": True},
            "validation": {"max_items": 10},
            "ui": {
                "placeholder": "搜索并选择人员",
                "multiple": True,
                "readonly": False,
                "hidden": False,
                "col_span": None,
            },
            "children": [],
        },
        {
            "name": "details",
            "label": "费用明细",
            "required": True,
            "type": "array",
            "component": "detail-table",
            "value_type": "array",
            "erp_field_type": "detail",
            "options": [],
            "option_values": [],
            "option_source": None,
            "validation": {"min_items": 1},
            "ui": {
                "placeholder": "至少添加一条明细",
                "multiple": True,
                "readonly": False,
                "hidden": False,
                "col_span": None,
            },
            "children": [
                {
                    "name": "item_name",
                    "label": "项目名称",
                    "required": False,
                    "type": "text",
                    "component": "text",
                    "value_type": "string",
                    "validation": {},
                    "ui": {"placeholder": "", "multiple": False, "readonly": False, "hidden": False, "col_span": None},
                    "children": [],
                },
                {
                    "name": "amount",
                    "label": "金额",
                    "required": False,
                    "type": "number",
                    "component": "money",
                    "value_type": "number",
                    "validation": {"min": 0, "scale": 2},
                    "ui": {"placeholder": "", "multiple": False, "readonly": False, "hidden": False, "col_span": None},
                    "children": [],
                },
            ],
        },
    ],
    "values": {
        "reason": "客户要求月底前完成交付",
        "expense_type": "travel",
        "notify_users": ["863"],
        "details": [{"item_name": "交通费", "amount": 320.50}],
    },
    "missing_field_keys": [],
    "invalid_fields": [],
}))
add_body(doc, "字段核心属性：")
add_table(doc, ["字段", "前端用途", "注意事项"], [
    ["name", "表单值对象的 key", "提交必须使用 name，不使用 label"],
    ["label", "展示标题", "只用于页面文案"],
    ["component", "选择前端组件", "未知类型不要默认当文本提交"],
    ["value_type", "决定值的 JavaScript 类型", "number、array、object 要保持类型"],
    ["required", "必填标记", "前后端都要校验"],
    ["option_values", "静态选项", "提交 value，不提交 label"],
    ["option_source", "远程选项来源", "lazy=true 时按需调用 options 接口"],
    ["validation", "长度、范围、精度、正则等", "前端提示，后端最终裁决"],
    ["ui", "隐藏、只读、占位符、多选等", "hidden 字段通常不要提交"],
    ["children", "明细表列或嵌套字段", "detail-table 按 children 动态生成列"],
], widths=[3.2, 5.7, 7.2])

add_heading(doc, "4 控件和提交值映射", 1)
add_table(doc, ["component", "前端控件", "提交值示例", "对应 value_type"], [
    ["text", "单行输入框", '"客户要求月底交付"', "string"],
    ["textarea", "多行文本框", '"详细说明内容"', "string"],
    ["number", "数字输入框", "10", "number"],
    ["money", "金额输入框", "3200.50", "number"],
    ["date", "日期选择器", '"2026-09-04"', "date"],
    ["datetime", "日期时间选择器", '"2026-09-04 14:30:00"', "datetime"],
    ["select", "下拉单选", '"travel"', "string"],
    ["radio", "单选按钮", '"urgent"', "string"],
    ["checkbox-group", "多选框", '["hotel", "meal"]', "array"],
    ["entity-select", "用户或部门选择器", '["863", "901"]', "array"],
    ["related-select", "订单、合同等关联对象", '["SO10001"]', "array"],
    ["approval-select", "ERP 字段中的审批关联选择", '["22"]', "array"],
    ["detail-table", "明细表", '[{"item_name":"交通费","amount":320.5}]', "array"],
    ["attachment", "附件控件", '[{"file_id":"file_1001"}]', "array"],
    ["address", "地址结构化控件", '{"province":"广东"}', "object"],
], widths=[3.2, 5.0, 6.0, 2.0])
add_body(doc, "单选、多选和输入框的最小字段示例：")
add_code(doc, json_text({
    "fields": [
        {
            "name": "leave_type",
            "label": "请假类型",
            "component": "select",
            "value_type": "string",
            "required": True,
            "option_values": [
                {"label": "年假", "value": "annual"},
                {"label": "事假", "value": "personal"},
            ],
        },
        {
            "name": "urgency",
            "label": "是否紧急",
            "component": "radio",
            "value_type": "string",
            "required": True,
            "option_values": [
                {"label": "普通", "value": "normal"},
                {"label": "紧急", "value": "urgent"},
            ],
        },
        {
            "name": "expense_items",
            "label": "费用项目",
            "component": "checkbox-group",
            "value_type": "array",
            "required": True,
            "ui": {"multiple": True},
            "option_values": [
                {"label": "交通费", "value": "transport"},
                {"label": "住宿费", "value": "hotel"},
                {"label": "餐费", "value": "meal"},
            ],
        },
        {
            "name": "remark",
            "label": "备注",
            "component": "textarea",
            "value_type": "string",
            "required": False,
            "validation": {"max_length": 500},
        },
    ],
    "values": {
        "leave_type": "annual",
        "urgency": "normal",
        "expense_items": ["transport", "hotel"],
        "remark": "",
    },
}))

add_heading(doc, "5 动态选项接口", 1)
add_body(doc, "当 option_source.lazy 为 true 时，前端在用户聚焦或输入关键字后调用 POST /api/approval/options。当前请求字段包括 template_id、field_key、keyword、page 和 page_size，并带上当前用户身份。")
add_code(doc, json_text({
    "user_id": "863",
    "company_id": "16",
    "template_id": "101",
    "field_key": "notify_users",
    "keyword": "张",
    "page": 1,
    "page_size": 20,
}))
add_code(doc, json_text({
    "template_id": "101",
    "field_key": "notify_users",
    "source": {"type": "user_list", "lazy": True, "searchable": True},
    "options": [
        {"label": "张三", "value": "863", "disabled": False, "meta": {"department": "技术部"}},
        {"label": "张四", "value": "901", "disabled": False, "meta": {"department": "财务部"}},
    ],
    "page": 1,
    "page_size": 20,
    "total": 2,
    "has_more": False,
}))
add_body(doc, "静态下拉、假期规则、人员、部门和关联业务对象都使用同一个 options 接口。前端只需要根据 field.option_source.type 展示不同的选择器；不要在浏览器中直接拼接 ERP 内部接口。")
add_table(doc, ["option_source.type", "典型字段", "页面行为"], [
    ["static", "普通下拉、单选、多选", "直接使用 option_values，可本地过滤"],
    ["holiday_rule", "假期规则", "服务端预置选项，通常不需要远程搜索"],
    ["user_list", "用户、员工", "防抖搜索，提交用户 ID"],
    ["related_list", "部门、订单、合同", "防抖搜索和分页，提交业务对象 ID"],
], widths=[4.2, 5.5, 6.5])

add_heading(doc, "6 多轮填写和校验返回", 1)
add_body(doc, "用户第一次输入业务意图后，服务端可能只返回模板或字段，前端应保留当前表单并继续让用户填写。每次补填都使用同一个 session_id，并把本轮值放在 form_values。")
waiting_user = {
    "message": "请补充或修正以下审批信息：缺少：leave_type；需要修正：结束时间必须晚于开始时间",
    "route": "approval_workflow",
    "assistant_type": "approval",
    "plan": {},
    "tool_calls": [],
    "errors": [],
    "pending_question": "请补充或修正以下审批信息：缺少：leave_type；需要修正：结束时间必须晚于开始时间",
    "workflow_status": "waiting_user",
    "erp_mode": "remote",
    "erp_write_mode": "disabled",
    "evidence": [],
    "citations": [],
    "erp_data": {},
    "form_schema": {
        "schema_version": "1.0",
        "template": {"template_id": "101", "template_code": "101", "title": "请假审批", "company_id": "16"},
        "fields": [
            {"name": "leave_type", "label": "请假类型", "component": "select", "value_type": "string", "required": True, "option_values": [{"label": "年假", "value": "annual"}, {"label": "事假", "value": "personal"}]},
            {"name": "start_time", "label": "开始时间", "component": "datetime", "value_type": "datetime", "required": True, "option_values": [], "validation": {}},
            {"name": "end_time", "label": "结束时间", "component": "datetime", "value_type": "datetime", "required": True, "option_values": [], "validation": {}},
        ],
        "values": {"start_time": "2026-09-05 09:00:00", "end_time": "2026-09-04 18:00:00"},
        "missing_field_keys": ["leave_type"],
        "invalid_fields": [{"field_key": "end_time", "message": "结束时间必须晚于开始时间"}],
    },
    "preview": None,
}
add_code(doc, json_text(waiting_user))
add_body(doc, "前端定位错误时使用 invalid_fields[].field_key 查找 fields[].name；不要只展示 pending_question 而不定位控件。")
add_body(doc, "补填请求示例：")
add_code(doc, json_text({
    "message": "补充请假类型为年假，并修正结束时间",
    "session_id": "approval-20260904-001",
    "request_id": "req-20260904-002",
    "assistant_key": "approval-assistant",
    "user_id": "863",
    "form_values": {
        "leave_type": "annual",
        "end_time": "2026-09-05 18:00:00",
    },
    "selected_assignees": {},
}))

add_heading(doc, "7 审批人选择和预览", 1)
add_body(doc, "字段校验通过后，服务端读取 ERP 审批节点并在 preview.approval_flow 返回候选审批人。前端不应自行查询或填写任意 UID，只能使用 candidates 中的人员。")
waiting_assignee = {
    "message": "表单字段已补齐，请在审批流程中选择审批人后再确认提交。",
    "route": "approval_workflow",
    "assistant_type": "approval",
    "plan": {},
    "tool_calls": [],
    "errors": [],
    "pending_question": "表单字段已补齐，请在审批流程中选择审批人后再确认提交。",
    "workflow_status": "waiting_assignee",
    "erp_mode": "remote",
    "erp_write_mode": "remote",
    "evidence": [],
    "citations": [],
    "erp_data": {},
    "form_schema": {"schema_version": "1.0", "template": {"template_id": "101", "template_code": "101", "title": "请假审批", "company_id": "16"}, "fields": [], "values": {}, "missing_field_keys": [], "invalid_fields": []},
    "preview": {
        "preview_id": "preview-10001",
        "preview_version": 1,
        "preview_hash": "a8c1f2d8...",
        "template_code": "101",
        "template_id": "101",
        "title": "请假审批",
        "fields": {"leave_type": "annual", "start_time": "2026-09-05 09:00:00", "end_time": "2026-09-05 18:00:00"},
        "submission_fields": {"leave_type": "annual", "start_time": "2026-09-05 09:00:00", "end_time": "2026-09-05 18:00:00"},
        "nodes": [],
        "submit_nodes": [],
        "approval_flow": [
            {
                "node_id": "22",
                "name": "部门负责人审批",
                "node_type": "approval",
                "level": 1,
                "handle_type": "submitter_choice",
                "requires_selection": True,
                "multiple": False,
                "candidates": [{"uid": "863", "name": "张三", "avatar": None}],
                "selected": [],
            }
        ],
        "selected_assignees": {},
        "missing_assignee_node_ids": ["22"],
        "invalid_assignee_nodes": [],
        "form_schema": {},
        "requires_confirmation": False,
        "idempotency_key": "idem-10001",
    },
}
add_code(doc, json_text(waiting_assignee))
add_body(doc, "选择审批人后，再次调用 /api/chat：")
add_code(doc, json_text({
    "message": "已选择张三，重新生成审批预览",
    "session_id": "approval-20260904-001",
    "request_id": "req-20260904-003",
    "assistant_key": "approval-assistant",
    "user_id": "863",
    "form_values": {},
    "selected_assignees": {"22": ["863"]},
}))
add_body(doc, "selected_assignees 的 key 是审批节点 node_id，value 始终是 UID 数组。即使节点只能单选，也仍使用数组传输。")

add_heading(doc, "8 确认提交和成功返回", 1)
add_body(doc, "当所有字段和审批人都满足要求时，服务端返回 workflow_status=preview_ready，并将 preview.requires_confirmation 设为 true。用户修改任何字段后都必须重新生成预览，旧 preview_id、preview_version 和 preview_hash 不可继续使用。")
preview_ready = {
    "message": "字段已补齐，已生成审批预览。请回复“确认提交”或“取消”。",
    "route": "approval_workflow",
    "assistant_type": "approval",
    "plan": {},
    "tool_calls": [{"tool": "erp.validate_fields", "valid": True, "field_count": 3, "node_count": 1}],
    "errors": [],
    "pending_question": "",
    "workflow_status": "preview_ready",
    "erp_mode": "remote",
    "erp_write_mode": "remote",
    "evidence": [],
    "citations": [],
    "erp_data": {},
    "form_schema": {},
    "preview": {
        "preview_id": "preview-10001",
        "preview_version": 2,
        "preview_hash": "f99a12...",
        "template_code": "101",
        "template_id": "101",
        "title": "请假审批",
        "submission_fields": {"leave_type": "annual", "start_time": "2026-09-05 09:00:00", "end_time": "2026-09-05 18:00:00"},
        "approval_flow": [{"node_id": "22", "name": "部门负责人审批", "multiple": False, "candidates": [{"uid": "863", "name": "张三", "avatar": None}], "selected": [{"uid": "863", "name": "张三", "avatar": None}]}],
        "selected_assignees": {"22": ["863"]},
        "missing_assignee_node_ids": [],
        "invalid_assignee_nodes": [],
        "requires_confirmation": True,
        "idempotency_key": "idem-10001",
    },
}
add_code(doc, json_text(preview_ready))
add_body(doc, "确认提交请求：")
add_code(doc, json_text({
    "message": "确认提交",
    "session_id": "approval-20260904-001",
    "request_id": "req-20260904-004",
    "assistant_key": "approval-assistant",
    "user_id": "863",
    "confirm": True,
    "preview_id": "preview-10001",
    "preview_version": 2,
    "preview_hash": "f99a12...",
    "form_values": {},
    "selected_assignees": {"22": ["863"]},
}))
submitted = {
    "message": "已完成提交：A202609040001。",
    "route": "approval_workflow",
    "assistant_type": "approval",
    "plan": {},
    "tool_calls": [{"tool": "erp.approval_submit", "mode": "remote"}],
    "errors": [],
    "pending_question": "",
    "workflow_status": "submitted",
    "erp_mode": "remote",
    "erp_write_mode": "remote",
    "evidence": [],
    "citations": [],
    "erp_data": {"approval_id": "A202609040001", "status": "已提交", "template_code": "101", "idempotency_key": "idem-10001"},
    "form_schema": None,
    "preview": None,
}
add_code(doc, json_text(submitted))
add_body(doc, "当 erp_write_mode=disabled 时，确认请求不会真实写入 ERP，通常返回 workflow_status=blocked，并在 erp_data 中说明“已阻止写入 ERP”。前端应展示明确提示，不要自动无限重试。")

add_heading(doc, "9 前端动态渲染建议", 1)
add_body(doc, "前端可以用一个字段渲染器覆盖所有审批类型。字段值使用 Record<string, unknown>，不要为每一种审批单独定义固定表单对象。")
add_code(doc, """type FormField = {
  name: string;
  label: string;
  component: string;
  value_type: 'string' | 'number' | 'date' | 'datetime' | 'array' | 'object';
  required?: boolean;
  option_values?: Array<{ label: string; value: unknown; disabled?: boolean; meta?: Record<string, unknown> }>;
  option_source?: { type: string; lazy?: boolean; searchable?: boolean; [key: string]: unknown } | null;
  validation?: Record<string, unknown>;
  ui?: { placeholder?: string; multiple?: boolean; readonly?: boolean; hidden?: boolean; [key: string]: unknown };
  children?: FormField[];
};

function renderField(field: FormField) {
  if (field.ui?.hidden) return null;

  const Component = componentMap[field.component];
  if (!Component) return <UnsupportedField field={field} />;

  return (
    <FormItem
      name={field.name}
      label={field.label}
      required={field.required}
      rules={buildRules(field)}
    >
      <Component
        valueType={field.value_type}
        options={field.option_values ?? []}
        optionSource={field.option_source}
        multiple={field.ui?.multiple}
        readonly={field.ui?.readonly}
        placeholder={field.ui?.placeholder}
        children={field.children}
      />
    </FormItem>
  );
}""", "ts")
add_body(doc, "组件注册表至少应覆盖 text、textarea、number、money、date、datetime、select、radio、checkbox-group、entity-select、related-select、approval-select、detail-table 和 attachment。未知 component 必须显示不支持提示，并阻止确认提交。")

add_heading(doc, "10 校验和字段值规则", 1)
add_table(doc, ["规则", "前端处理", "后端处理"], [
    ["required", "空值时在字段下提示必填", "再次校验并返回 missing_field_keys"],
    ["min / max", "数字或金额范围校验", "再次校验，返回 invalid_fields"],
    ["min_length / max_length", "输入长度校验", "再次校验"],
    ["scale / precision", "金额保留小数位提示或限制", "最终按 ERP 规则校验"],
    ["pattern", "正则校验", "再次校验；非法正则不应让页面崩溃"],
    ["min_items / max_items", "多选和明细表数量校验", "再次校验"],
    ["时间顺序", "开始时间不能晚于结束时间", "服务端执行跨字段校验"],
    ["枚举值", "只能提交 option_values 中的 value", "服务端检查字段白名单和选项合法性"],
], widths=[4.4, 6.0, 6.1])
add_body(doc, "提交值必须保持正确类型：金额和数字传 number，多选和人员传 array，地址等结构化字段传 object。下拉框和单选传 option_values[].value，不要传展示文字。")
add_code(doc, json_text({
    "form_values": {
        "amount": 3200.5,
        "expense_type": "travel",
        "expense_items": ["transport", "hotel"],
        "notify_users": ["863", "901"],
        "work_address": {"province": "广东省", "city": "深圳市", "detail": "科技园"},
    }
}))

add_heading(doc, "11 历史记录和页面状态", 1)
add_body(doc, "左侧会话列表和右侧消息区使用 session_id 关联。切换会话后要同时恢复聊天文字、form_schema、values、preview 和 workflow_status，不能只恢复 message。")
add_code(doc, json_text({
    "company_id": "16",
    "user_id": "863",
    "assistant_key": "approval-assistant",
    "status": "active",
    "page": 1,
    "page_size": 20,
}))
add_code(doc, json_text({
    "items": [
        {
            "session_id": "approval-20260904-001",
            "title": "请假审批",
            "status": "active",
            "workflow_status": "preview_ready",
            "last_message_seq": 6,
            "last_active_at": "2026-09-04T11:30:00",
        }
    ],
    "count": 1,
    "page": 1,
    "page_size": 20,
    "has_more": False,
}))
add_body(doc, "页面状态建议：")
add_code(doc, "idle -> collecting_fields -> waiting_user -> waiting_assignee -> preview_ready -> submitted\n                              |                         |\n                              +-> waiting_erp            +-> blocked / failed / cancelled", "text")
add_body(doc, "只有 workflow_status=preview_ready 且 preview.requires_confirmation=true 时显示确认提交按钮。submitted、blocked、cancelled 和 failed 状态应结束当前草稿，避免重复确认。")

add_heading(doc, "12 错误处理和落地检查清单", 1)
add_table(doc, ["检查项", "实现要求"], [
    ["响应解析", "直接读取 message、form_schema、preview、erp_data，不假设存在 content、answer、data"],
    ["控件选择", "只根据 component 选择组件；未知类型阻止提交"],
    ["字段定位", "用 field.name 与 missing_field_keys、invalid_fields[].field_key 对应"],
    ["选项值", "提交 value；label 只用于展示"],
    ["远程选项", "防抖、分页、取消过期请求，并保留 has_more"],
    ["审批人", "只允许 candidates 中的 UID；单选节点最多一个"],
    ["预览", "字段或审批人变化后重新生成 preview"],
    ["重复提交", "禁用按钮，并保留 request_id 和 idempotency_key"],
    ["附件", "先通过 ERP 上传接口拿到文件 ID，再提交引用；不能提交本地路径"],
    ["安全", "不要在前端保存 LLM、Milvus、MySQL 或 ERP 密钥"],
], widths=[4.5, 12.0])
add_body(doc, "目前需要后续补充 ERP 元数据或接口的能力包括：条件必填、字段显隐联动、复杂明细表公式、统一附件上传与删除、字段级权限和标准化业务规则错误。前端可以先按本协议完成通用动态表单，后续新增 component 时只增加组件映射。")

doc.core_properties.title = "ERP 审批动态表单前端对接说明"
doc.core_properties.subject = "ERP 审批动态字段渲染与提交 JSON 契约"
doc.core_properties.author = "ERP AI Assistant"
doc.core_properties.comments = ""
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
