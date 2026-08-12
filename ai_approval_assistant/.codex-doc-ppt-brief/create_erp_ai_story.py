from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\PythonProject\LearnOne\ERP与AI融合实践及体验提升路线.docx")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(0, 121, 107)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
GRAY = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4D6"
LIGHT_RED = "FDECEC"


def set_run(run, size=11, color=None, bold=False, italic=False):
    font = "Microsoft YaHei"
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), font)
    rpr.rFonts.set(qn("w:hAnsi"), font)
    rpr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, color=None, bold=False, italic=False, before=0, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    if text:
        set_run(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 11 if level == 2 else 8)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    set_run(
        paragraph.add_run(text),
        size=17 if level == 1 else 13 if level == 2 else 11.5,
        color=BLUE if level < 3 else NAVY,
        bold=True,
    )
    return paragraph


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(text), size=10.8)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(text), size=10.8)
    return paragraph


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 100), ("bottom", 100), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_callout(doc, label, content, fill=LIGHT_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(f"{label}："), size=10.8, color=label_color, bold=True)
    set_run(paragraph.add_run(content), size=10.8)
    add_para(doc, after=2)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(paragraph.add_run(header), size=9.6, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            paragraph = row.cells[index].paragraphs[0]
            set_run(paragraph.add_run(value), size=9.3)
    set_table_geometry(table, widths)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.82)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("ERP 与 AI 融合实践及体验提升路线"), size=9, color=GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

add_para(doc, "AI 效能分享", size=12, color=BLUE, bold=True, before=28, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "从复杂表单到智能业务入口", size=28, color=NAVY, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "ERP 与 AI 融合实践及体验提升路线", size=17, color=BLUE, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "以审批助手、日报助手和四种 Agent 实现为实践基础", size=11, color=GRAY, after=38, align=WD_ALIGN_PARAGRAPH.CENTER)
add_callout(
    doc,
    "核心观点",
    "我们的目标不是把 ERP 做成聊天机器人，也不是让大模型替代业务系统，而是在 ERP 之上增加一层自然语言交互和受控任务执行能力，让员工从‘寻找系统功能’转变为‘表达业务目标’。",
)
add_para(doc, "本稿将当前四种实现放回 ERP 体验提升的主线中：它们不是孤立的技术实验，而是探索 AI 如何逐步、安全地进入 ERP 业务流程。", size=10.5, color=GRAY, italic=True, after=18)

add_heading(doc, "一、我们真正要解决的不是表单，而是 ERP 的使用门槛", 1)
add_para(doc, "ERP 擅长沉淀数据、权限和流程，但对普通员工而言，完成一个业务目标往往需要先理解系统。", size=11.5, color=NAVY, bold=True)
for item in [
    "员工需要知道功能入口在哪里、应该选择哪个模板、字段如何填写。",
    "员工需要理解审批规则、关联数据和后续操作，而这些知识通常分散在页面、文档和个人经验中。",
    "日报、审批、采购等场景存在大量重复查找、重复填写和重复整理。",
    "新员工或低频用户需要反复学习系统，ERP 功能越丰富，使用成本也可能越高。",
]:
    add_bullet(doc, item)
add_para(doc, "但用户真正关心的并不是菜单和字段，而是业务目标：", size=11, bold=True, before=5)
add_callout(doc, "用户表达", "我要报销去上海拜访客户产生的差旅费。\n帮我整理今天的工作并生成日报。\n我要采购 10 台笔记本电脑，预算 6 万元。", LIGHT_GRAY, NAVY)
add_para(doc, "因此，ERP 与 AI 结合的第一价值不是替代现有流程，而是改变人与 ERP 的交互方式：", size=11, bold=True)
add_matrix(
    doc,
    ["过去", "未来"],
    [
        ("人适应系统、寻找入口、理解字段", "人描述目标，AI 组织系统能力"),
        ("用户重复查找和录入数据", "AI 从 ERP 加载已有数据并补齐信息"),
        ("提交后才发现字段或规则错误", "提交前完成校验、预览和人工确认"),
    ],
    [4680, 4680],
)

add_heading(doc, "二、目标形态：ERP 仍是核心，AI 成为新的业务交互层", 1)
add_para(doc, "企业 AI 不能理解成“大模型直接操作 ERP”。审批、采购、库存和数据修改都涉及真实权限、组织规则和写入风险，必须采用分层协作。", size=11.2)
add_callout(
    doc,
    "分工原则",
    "大模型负责理解用户，Workflow 负责遵守公司规则，Tools 负责调用系统能力，ERP 负责真实数据和执行结果，用户负责高风险操作的最终授权。",
    LIGHT_GOLD,
    GOLD,
)
add_matrix(
    doc,
    ["层次", "主要职责", "不能承担的职责"],
    [
        ("大模型 / Agent", "意图理解、字段提取、内容整理、任务规划", "不能虚构模板、权限、审批人和提交结果"),
        ("Workflow", "流程状态、必经步骤、业务 Gate、中断恢复", "不负责替代 ERP 业务事实"),
        ("ERP Tools", "将查询、保存、预览、提交封装成标准能力", "不能绕过权限、校验和审计"),
        ("ERP", "真实数据、模板、权限、流程和执行结果", "不需要承担自然语言理解"),
        ("人工确认", "批准高风险写操作", "不应承担重复字段填写和系统查找"),
    ],
    [1600, 3880, 3880],
)

add_heading(doc, "三、当前实践：四种实现其实是在探索四种执行权配置", 1)
add_para(doc, "审批助手和日报助手的四种写法看起来有重复，但这些重复有价值：同一目标反复实现，才能看清哪些能力适合交给模型，哪些边界必须保留在代码和 ERP 中。", size=11.2)

add_heading(doc, "1. 固定 Flow：daily_report_agent", 2)
add_para(doc, "固定 Flow 将业务步骤直接写进 LangGraph：动作识别、日期确认、上下文加载、内容收集、草稿保存、预览、确认、提交、修改和取消都有明确节点。")
for item in [
    "证明了日报业务可以通过自然语言入口完成完整闭环。",
    "证明了状态、预览和人工确认可以被代码稳定控制。",
    "适合正式日报、审批、支付和数据修改等强规则、高风险业务。",
    "问题是规则越多，节点、条件边和修改分支会越来越复杂。",
]:
    add_bullet(doc, item)

add_heading(doc, "2. 自主 Agent：daily_report_create_agent", 2)
add_para(doc, "自主 Agent 将日期、上下文、草稿和预览等工具交给模型，由模型决定下一步调用什么。提交工具仍保留 confirmed=true 守卫。")
for item in [
    "证明了模型可以围绕业务目标自主组合 ERP 工具。",
    "减少了开发者提前写死每条工具调用路径的需要。",
    "问题是业务阶段需要从 messages 和 tool calls 中反推。",
    "工具顺序、循环次数、Token、响应时间和结果稳定性更难预测。",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Agentic Workflow：daily_report_agentic_workflow_demo", 2)
add_para(doc, "Agentic Workflow 将理解和执行拆开：Agent 提取日期与工作内容、生成计划和正文；Workflow 确认权威日期、加载 ERP 上下文、保存草稿、生成预览并控制提交 Gate。")
for item in [
    "证明了 Agent 的灵活性和 Workflow 的确定性可以结合。",
    "计划、上下文、生成结果和 Gate 状态分开保存，更容易观察和测试。",
    "适合智能填单、报告生成和其他既需要自然语言理解又需要安全写入的场景。",
    "当前规划和生成节点仍以确定性规则模拟，保存和提交使用 Demo 状态。",
]:
    add_bullet(doc, item)

add_heading(doc, "4. Deep Agent：ai_deep_agents_assistant", 2)
add_para(doc, "Deep Agent 在审批场景中探索目标驱动、多工具调用、MemorySaver checkpoint、interrupt_on 和人工恢复。")
for item in [
    "证明了开放任务可以由 Agent 根据目标自主选择用户、模板、草稿、预览和提交工具。",
    "证明了危险写操作可以在工具执行前中断，并在用户批准后从 checkpoint 恢复。",
    "适合跨多个 ERP 模块、执行路径依赖中间结果的长任务。",
    "当前业务后端是 Mock 审批服务，不能作为生产审批已落地的证明。",
]:
    add_bullet(doc, item)

add_heading(doc, "四、当前我们走到哪里：从单点助手迈向受控业务执行层", 1)
add_callout(
    doc,
    "阶段判断",
    "当前已经完成‘自然语言入口 + 单场景业务闭环’的验证，正在从单点 AI 助手迈向‘受控 ERP 业务执行层’。我们还没有进入公司级跨模块智能体阶段。",
)
add_matrix(
    doc,
    ["能力阶段", "目标", "当前状态"],
    [
        ("阶段 1：自然语言入口", "用户不再找菜单，直接表达业务目标", "已验证：审批、日报意图可以通过问答进入"),
        ("阶段 2：智能填单与预览", "提取字段、补齐信息、生成业务预览", "已基本验证：日报完整；审批在 Demo/Mock 中验证"),
        ("阶段 3：受控业务执行", "真实权限、校验、Gate、人工确认和写入", "正在建设：部分真实工具已接入，统一边界尚未完成"),
        ("阶段 4：ERP 能力平台化", "将各模块接口沉淀为标准、可复用工具", "尚未系统化：当前工具仍围绕单个场景建设"),
        ("阶段 5：跨模块 Agent", "根据目标动态组合客户、项目、采购、库存等能力", "探索阶段：Deep Agent 验证架构，不具备生产业务闭环"),
    ],
    [2200, 3300, 3860],
)
add_para(doc, "因此，现阶段最重要的工作不是继续增加更多 Agent 名称，而是把已有场景向真实 ERP 能力、统一状态和安全执行推进。", size=11.2, color=NAVY, bold=True, before=8)

add_heading(doc, "五、审批与日报分别证明了什么", 1)
add_heading(doc, "审批：证明 AI 可以降低复杂表单的操作门槛", 2)
for item in [
    "用户只需要描述报销、请假或采购目标。",
    "AI 可以辅助识别业务类型、提取金额、事由、日期和其他字段。",
    "ERP 必须提供真实模板、用户权限、审批节点和执行结果。",
    "提交前必须生成预览；没有明确确认不能执行写操作。",
]:
    add_bullet(doc, item)
add_callout(doc, "审批价值", "AI 将复杂 ERP 表单转化为自然语言交互，但不改变 ERP 原有的权限、规则和审批安全边界。", LIGHT_GRAY, NAVY)

add_heading(doc, "日报：证明 AI 可以复用 ERP 数据并减少重复录入", 2)
for item in [
    "从 ERP 加载日报日期、草稿、配置和可关联业务记录。",
    "用户只补充当天工作内容，AI 负责整理和生成规范文本。",
    "Workflow 负责保存草稿、生成预览和控制提交。",
    "未来可关联审批、客户、项目、采购和库存等当天业务记录。",
]:
    add_bullet(doc, item)
add_callout(doc, "日报价值", "ERP 提供真实工作数据，AI 负责理解和整理，Workflow 负责安全提交，从而减少员工回忆和重复填写。", LIGHT_GRAY, NAVY)

add_heading(doc, "六、ERP + AI 如何真正提升用户体验", 1)
add_para(doc, "体验提升不能只看聊天界面是否好看，而要看员工完成真实业务目标时减少了什么成本。", size=11.2, color=NAVY, bold=True)
add_matrix(
    doc,
    ["体验问题", "AI + ERP 改进方式", "预期价值"],
    [
        ("不知道功能在哪里", "统一自然语言业务入口，自动路由到对应能力", "降低菜单查找和培训成本"),
        ("不知道表单怎么填", "从自然语言提取字段，动态追问缺失信息", "降低理解成本和填写错误"),
        ("已有数据还要重复输入", "从 ERP 加载用户、组织、项目和业务记录", "减少重复录入，提高数据完整性"),
        ("提交前不知道结果", "生成结构化预览、审批节点和风险提示", "提升可预期性，减少误提交"),
        ("修改或取消容易混乱", "通过状态机区分修改、取消提交和取消流程", "提升多轮交互的一致性"),
        ("跨模块任务需要反复切换", "Agent 动态组合多个标准 ERP 工具", "减少系统切换和人工汇总"),
    ],
    [2300, 4000, 3060],
)

add_heading(doc, "七、下一步路线：先把单点闭环做实，再建设公司级 AI 业务能力层", 1)
add_heading(doc, "近期：把现有审批和日报做成可稳定演示、可持续验证的真实闭环", 2)
for item in [
    "统一会话状态、意图识别和路由规则，解决 session 串线和取消语义混淆。",
    "接入真实 ERP 用户、权限、模板、日报和审批接口，减少 Demo 与 Mock 依赖。",
    "统一预览结构、确认语句、错误返回和幂等机制。",
    "增加 trace、工具参数、耗时、Token 和失败原因的可观测性。",
    "为固定 Flow、Agent 和 Agentic Workflow 建立相同的业务验收指标。",
]:
    add_bullet(doc, item)

add_heading(doc, "中期：将 ERP 接口沉淀为标准业务工具", 2)
for item in [
    "用户与组织能力：当前用户、直属上级、部门和权限。",
    "审批能力：模板、动态字段、审批节点、草稿、预览和提交。",
    "日报能力：日期、配置、同步数据、草稿、预览和提交。",
    "客户、项目、采购、库存和消息提醒能力。",
    "所有工具统一输入输出、权限校验、日志审计、幂等和异常处理。",
]:
    add_bullet(doc, item)

add_heading(doc, "长期：形成跨模块 AI 业务能力层", 2)
for item in [
    "根据本周客户拜访记录，识别长期未跟进客户，生成周报并提醒负责人。",
    "查询项目预算、采购金额、入库情况和待审批申请，分析超预算风险。",
    "根据当天审批、客户、项目和任务记录自动整理工作总结。",
    "根据业务目标动态组合多个 ERP 工具，但所有高风险写操作仍需 Workflow Gate 和人工确认。",
]:
    add_bullet(doc, item)

add_heading(doc, "八、推荐的技术路线：业务风险决定 Agent 自主程度", 1)
add_matrix(
    doc,
    ["任务类型", "建议模式", "原因"],
    [
        ("规则明确、高风险写入", "固定 Flow / Workflow", "路径必须确定，状态、审计和回滚最重要"),
        ("需要智能理解，但写入有严格顺序", "Agentic Workflow", "Agent 处理不确定性，Workflow 控制业务 Gate"),
        ("开放查询、内容整理、低风险辅助", "自主 Agent", "允许工具顺序变化，追求灵活性"),
        ("跨模块、长任务、多工具动态组合", "Deep Agent + HITL", "路径难以提前枚举，需要 checkpoint 和人工中断"),
    ],
    [2500, 2700, 4160],
)
add_callout(doc, "路线原则", "任务越确定，Workflow 权重越高；任务越开放，Agent 自主性越高；操作风险越高，确定性校验、审计和人工确认越重要。", LIGHT_GOLD, GOLD)

add_heading(doc, "九、如何衡量 ERP 体验是否真的提升", 1)
add_para(doc, "建设成果不能用“调用了多少次大模型”衡量，而要回到真实业务闭环和用户体验。", size=11.2, color=NAVY, bold=True)
for item in [
    "完成一个业务目标需要多少页面、点击和对话轮次。",
    "用户首次完成任务的成功率，以及缺失字段的自动补齐率。",
    "是否接入真实 ERP 数据、权限和业务规则。",
    "预览后修改率、误提交率和人工撤回率是否下降。",
    "平均响应时间、模型调用次数、Token 成本和失败重试率。",
    "工具和流程是否可追踪、可审计、可复用。",
    "新员工培训时间和低频业务咨询数量是否下降。",
    "同一项 ERP 能力能否被多个 Workflow 和 Agent 复用。",
]:
    add_bullet(doc, item)

add_heading(doc, "十、建议演示主线：用四个案例证明 ERP 体验提升的四个层次", 1)
add_number(doc, "自然语言路由：通过‘我的上级是谁？’证明系统能够理解目标并选择用户与组织能力。")
add_number(doc, "复杂审批简化：通过差旅报销证明自然语言可以完成模板匹配、字段提取、预览和确认。")
add_number(doc, "减少重复录入：通过写日报证明 AI 可以复用 ERP 上下文、整理内容并保存预览。")
add_number(doc, "受控智能执行：通过 Agentic Workflow 和 Deep Agent 证明模型可以规划工具，但 Workflow Gate、checkpoint 和人工确认仍然控制风险。")
add_callout(
    doc,
    "演示收束",
    "自然语言降低操作门槛 → Workflow 固化业务规则 → Tools 沉淀 ERP 标准能力 → Agent 组合多项业务能力 → 形成公司级 AI 业务能力层。",
)

add_heading(doc, "十一、最终结论", 1)
add_para(doc, "我们的目标不是让 ERP 变成一个聊天机器人，也不是让大模型替代现有业务系统。", size=12, color=NAVY, bold=True)
add_para(doc, "我们希望通过 AI 降低员工使用 ERP 的门槛，让员工从“寻找系统功能”转变为“表达业务目标”。大模型负责理解和整理，Workflow 负责固化公司流程，Tools 负责封装 ERP 能力，ERP 负责真实数据和业务执行，高风险操作最终由用户确认。", size=11.2)
add_para(doc, "审批和日报只是第一批场景。当前已经证明自然语言入口和单场景闭环可行，下一步要把真实接口、权限、状态、审计和人工确认做成标准能力。随着用户、客户、项目、采购、库存等能力逐步工具化，这些能力可以被不同 Workflow 和 Agent 重复组合，形成公司的 AI 业务能力层。", size=11.2)
add_callout(
    doc,
    "一句话结论",
    "AI 负责理解，代码负责边界，ERP 负责事实；让 AI 获得合适的决策权，而不是最大的决策权。",
    LIGHT_GOLD,
    GOLD,
)

add_heading(doc, "附录：当前实现的真实性边界", 1)
for item in [
    "daily_report_agent：固定 LangGraph 日报子图，使用现有日报工具，适合展示真实固定流程。",
    "daily_report_create_agent：使用真实 Agent 和工具调用，需要模型 API，提交工具包含确认守卫。",
    "daily_report_agentic_workflow_demo：规划和生成节点为规则模拟；上下文加载使用现有工具；保存和提交为 Demo 状态。",
    "ai_deep_agents_assistant：使用真实 Deep Agents、MemorySaver 和 interrupt_on；业务审批后端为 Mock 服务。",
    "三个日报实验必须使用不同 session_id；同一流程确认提交必须继续使用原 session_id。",
]:
    add_bullet(doc, item)

doc.core_properties.title = "ERP与AI融合实践及体验提升路线"
doc.core_properties.subject = "审批助手、日报助手及ERP AI业务能力层建设路线"
doc.core_properties.author = ""

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)

with ZipFile(OUTPUT) as archive:
    error = archive.testzip()
if error:
    raise RuntimeError(error)

check = Document(OUTPUT)
print({
    "output": str(OUTPUT),
    "size": OUTPUT.stat().st_size,
    "paragraphs": len(check.paragraphs),
    "tables": len(check.tables),
    "headings": len([p for p in check.paragraphs if p.style.name.startswith("Heading")]),
})
