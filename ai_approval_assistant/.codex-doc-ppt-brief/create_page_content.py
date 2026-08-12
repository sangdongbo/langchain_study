from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\PythonProject\LearnOne\AI效能分享-审批与日报四种AI实现-逐页内容稿.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(89, 89, 89)


def set_run(run, size=11, color=None, bold=False):
    font = "Microsoft YaHei"
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def paragraph(doc, text="", size=11, color=None, bold=False, after=6, before=0, alignment=None):
    item = doc.add_paragraph()
    item.paragraph_format.space_before = Pt(before)
    item.paragraph_format.space_after = Pt(after)
    item.paragraph_format.line_spacing = 1.2
    if alignment is not None:
        item.alignment = alignment
    if text:
        set_run(item.add_run(text), size=size, color=color, bold=bold)
    return item


def bullet(doc, text):
    item = doc.add_paragraph(style="List Bullet")
    item.paragraph_format.space_after = Pt(4)
    item.paragraph_format.line_spacing = 1.15
    set_run(item.add_run(text), size=11)
    return item


def number(doc, text):
    item = doc.add_paragraph(style="List Number")
    item.paragraph_format.space_after = Pt(4)
    item.paragraph_format.line_spacing = 1.15
    set_run(item.add_run(text), size=10.5)
    return item


def heading(doc, text, level=1, page_break=False):
    item = doc.add_paragraph(style=f"Heading {level}")
    item.paragraph_format.page_break_before = page_break
    item.paragraph_format.keep_with_next = True
    item.paragraph_format.space_before = Pt(0 if page_break else 8)
    item.paragraph_format.space_after = Pt(6)
    set_run(item.add_run(text), size=17 if level == 1 else 12.5, color=BLUE if level == 1 else NAVY, bold=True)
    return item


def add_slide(doc, page, title, subtitle, content, speech, demo=None, note=None):
    heading(doc, f"第 {page} 页", 1, page_break=page != 1)
    paragraph(doc, f"页面标题：{title}", size=15, color=NAVY, bold=True, after=4)
    if subtitle:
        paragraph(doc, f"副标题：{subtitle}", size=11, color=GRAY, after=10)
    heading(doc, "页面内容", 2)
    for item in content:
        bullet(doc, item)
    if demo:
        heading(doc, "现场操作", 2)
        for item in demo:
            number(doc, item)
    if note:
        heading(doc, "必须说明", 2)
        paragraph(doc, note, size=10.5, color=GRAY)
    heading(doc, "这一页怎么讲", 2)
    paragraph(doc, speech, size=10.5)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

for style_name in ["Heading 1", "Heading 2"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("AI 效能分享｜逐页内容稿"), size=9, color=GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

paragraph(doc, "AI 效能分享 PPT 逐页内容稿", size=24, color=NAVY, bold=True, before=18, after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
paragraph(doc, "从复杂表单到一句话提交", size=16, color=BLUE, bold=True, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
paragraph(doc, "审批助手与日报的四种 AI 实现", size=11, color=GRAY, after=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
paragraph(doc, "以下每个章节对应一页 PPT，可以直接复制页面标题和页面内容。‘这一页怎么讲’放在演讲者备注里，不需要显示在 PPT 页面上。", size=10.5, color=GRAY, after=20)

add_slide(doc, 1, "从复杂表单到一句话提交", "审批助手与日报的四种 AI 实现", [
    "固定 Flow",
    "自主 Agent",
    "Agentic Workflow",
    "Deep Agent",
], "今天分享的不是一个已经全面上线的产品，而是我围绕审批和日报场景进行的四次 AI 实现探索。目标一直没有变化：用户通过一句自然语言完成复杂信息的收集、预览和提交。")

add_slide(doc, 2, "今天分享什么", "从业务问题到四种实现，再到现场演示", [
    "应用场景：为什么要做问答式审批助手",
    "实施路径：从固定 Flow 到 Agent，再到混合架构",
    "现场演示：三种日报实现和一个 Deep Agent 审批实现",
    "预期效果：当前已经完成哪些功能",
    "遇到的问题：状态、边界和演示稳定性",
], "这五部分正好覆盖通知要求的应用场景、实施路径、预期效果和遇到的问题。为了让差异更直观，中间会连续演示四种实现。")

add_slide(doc, 3, "应用场景：平台审批为什么需要 AI", "用户不应该先学会表单，才能完成审批", [
    "当前操作：寻找系统入口、选择审批模板、理解字段含义、逐项填写、确认审批人、提交。",
    "用户真正想表达的通常很简单：我要请假三天、我要报销差旅费、帮我写今天的日报。",
    "希望的体验：用户直接描述目标，系统自动理解意图、追问缺失信息、生成业务预览。",
    "最终提交前仍然由用户进行明确确认。",
], "这个功能不是为了炫技，而是为了解决入口多、字段多、表单难理解的问题。AI 的价值是理解自然语言，业务系统仍然负责字段校验、流程规则和最终写入。")

add_slide(doc, 4, "实施路径：从 Flow 到 Deep Agent", "每次演进都在重新分配谁来做决定", [
    "第一阶段：固定 Flow。开发者定义节点、条件边和执行顺序。",
    "第二阶段：自主 Agent。模型自己决定下一步和调用什么工具。",
    "第三阶段：Agentic Workflow。Agent 负责理解和生成，Workflow 负责业务规则和提交 Gate。",
    "第四阶段：Deep Agent。探索复杂目标、Checkpoint、线程恢复和 Human-in-the-loop。",
], "四种实现不是简单的版本升级，也不是越新的越好。它们最大的区别是执行权交给谁，以及业务风险由谁控制。")

add_slide(doc, 5, "固定 Flow：daily_report_agent", "执行顺序由 Workflow 完全控制", [
    "识别用户动作：写日报、查看、提交、修改或取消。",
    "确认日报日期并加载日报上下文。",
    "收集工作内容并保存日报草稿。",
    "生成完整日报预览。",
    "等待用户确认后提交；也支持修改日期、修改内容或取消。",
    "优势：节点状态清晰、流程稳定、容易排查和编写测试。",
], "最开始我使用 LangGraph 的 Flow 模式，把每一个业务步骤都写成节点和条件边。它不一定最智能，但业务路径最清楚，适合规则明确的生产流程。")

add_slide(doc, 6, "现场演示 1：固定 Flow 写日报", "重点观察固定路径、预览和确认", [
    "提问：写日报：今天完成审批助手开发",
    "观察是否进入 daily_report_agent。",
    "观察 trace 中的节点是否按照固定顺序执行。",
    "观察系统是否先保存草稿、生成预览，再等待确认。",
    "回复“确认提交”后，流程才继续执行提交。",
], "演示时先发起写日报请求，再展示 trace 和日报预览，最后使用相同 session 确认提交。固定 Flow 的重点是过程稳定、问题可以直接定位到某个节点。", demo=[
    "调用 http://127.0.0.1:8010/docs 中的 POST /api/ai-approval/chat。",
    "使用 session_id：demo-flow-001。",
    "发送：写日报：今天完成审批助手开发。",
    "展示 agent、status、preview 和 trace。",
    "继续使用 demo-flow-001 发送：确认提交。",
])

add_slide(doc, 7, "自主 Agent：daily_report_create_agent", "模型自己选择工具和下一步", [
    "使用 create_agent；不可用时回退 create_react_agent。",
    "Agent 可以自主调用当前日期、加载日报上下文、保存草稿、生成预览等工具。",
    "每次工具调用后，模型观察结果，再决定继续调用工具还是回答用户。",
    "提交工具增加 confirmed=true 守卫，未明确确认时不能提交。",
    "状态主要保存在 daily_report_agent_messages 中。",
    "问题：工具顺序、循环次数、Token、延迟和成本更难预测。",
], "第二种方式想解决固定 Flow 节点越来越多的问题，所以尝试把工具交给 Agent 自己选择。但自由度增加后，业务阶段需要从消息和 tool calls 中反推，排查成本也增加了。")

add_slide(doc, 8, "现场演示 2：自主 Agent 写日报", "重点观察 messages 和工具调用顺序", [
    "提问：自主agent写日报：今天完成审批助手开发2",
    "观察 daily_report_agent_messages。",
    "观察模型调用了哪些日报工具。",
    "观察日期、上下文、草稿和预览工具的调用顺序。",
    "最终结果可能和固定 Flow 相同，但过程由 Agent 自己决定。",
], "这里不要只展示最后的日报结果。重点展示 Agent 消息历史和工具调用，让大家看到自主 Agent 与固定 Flow 的核心差异。", demo=[
    "使用新的 session_id：demo-agent-001。",
    "发送：自主agent写日报：今天完成审批助手开发2。",
    "查看 AgentMessage、ToolMessage 和 tool calls。",
    "展示模型实际选择的工具及参数。",
], note="该方式需要真实模型或 API。模型输出或工具顺序可能发生变化。")

add_slide(doc, 9, "Agentic Workflow：daily_report_agentic_workflow_demo", "Agent 负责智能判断，Workflow 负责业务规则", [
    "Agent 规划：理解用户输入，提取日期表达和工作内容，给出下一步建议。",
    "Workflow 确认日期并加载现有日报上下文。",
    "Agent 组织日报正文。",
    "Workflow 保存草稿并生成预览。",
    "Preview Gate：没有预览不能提交。",
    "Submit Gate：用户没有明确确认不能提交。",
    "中间状态分别保存 plan、context、compose、events 和 draft_saved。",
], "这套方式想保留 Agent 的理解和生成能力，同时让 Workflow 控制写入顺序。它不是依赖 Prompt 提醒模型不要直接提交，而是由代码中的 Gate 直接限制。", note="当前 demo_agent_plan 和 demo_agent_compose 是规则模拟；保存草稿和最终提交使用 Demo 状态。")

add_slide(doc, 10, "现场演示 3：Agentic Workflow 写日报", "重点观察职责分工、结构化状态和 Gate", [
    "提问：agentic workflow写日报：今天完成审批助手开发3",
    "观察 plan、context 和 compose 是否分别保存。",
    "观察 events 是否显示 Agent 和 Workflow 交替工作。",
    "观察 trace 是否进入 demo_preview_gate。",
    "没有预览、没有明确确认时都不能提交。",
    "确认后进入 demo_submit_gate。",
], "这次演示要突出为什么它比完全自主 Agent 更适合业务写入：Agent 提供智能空间，Workflow 提供稳定顺序，Gate 保证提交边界。", demo=[
    "使用新的 session_id：demo-agentic-001。",
    "发送：agentic workflow写日报：今天完成审批助手开发3。",
    "展示 plan、context、compose、events 和 trace。",
    "继续使用 demo-agentic-001 发送：确认提交。",
], note="最终返回的是 Demo 提交编号，不代表已经真实写入生产 ERP。")

add_slide(doc, 11, "三种日报实现对比", "同一目标，不同的执行权", [
    "固定 Flow：Workflow 控制顺序；节点状态最清晰；成本和调用次数更可控；适合稳定主流程。",
    "自主 Agent：Agent 自主选择工具；状态需要从消息和工具历史反推；适合工具调用实验。",
    "Agentic Workflow：Agent 建议、Workflow 决定；计划、上下文、生成和 Gate 分层；适合受控智能流程。",
    "选择原则：写入风险越高，代码控制应该越多；任务越开放、风险越低，Agent 权限可以越大。",
], "这一页不要把三种方式理解为强弱排名。实际选择要看业务确定性、写入风险、状态可观测性和成本预算。")

add_slide(doc, 12, "Deep Agent：ai_deep_agents_assistant", "探索复杂目标、Checkpoint 和人工确认", [
    "使用 create_deep_agent。",
    "使用 MemorySaver 保存线程 checkpoint。",
    "工具包括用户上下文、审批模板、审批草稿、审批预览和提交。",
    "submit_approval_request 配置 interrupt_on。",
    "危险提交前暂停，等待人工批准。",
    "确认后通过 Command(resume) 恢复原线程。",
    "当前支持请假、报销和采购等审批模板。",
], "Deep Agent 这次实现的是审批场景，不是日报，所以不能和前三种做完全公平的性能比较。它主要验证更长线程、多工具和 Human-in-the-loop。", note="Deep Agent 和 checkpoint 是真实实现，但业务提交后端是确定性的 Mock 服务。")

add_slide(doc, 13, "现场演示 4：Deep Agent 报销审批", "重点观察 Checkpoint、interrupt 和 resume", [
    "提问：我要报销差旅费，金额 5200 元，因为去上海拜访客户，发票已提供。",
    "第一轮：Agent 查询用户信息和审批模板，收集草稿并生成预览。",
    "危险提交工具执行前触发 interrupt。",
    "系统返回 awaiting_confirmation。",
    "使用相同 session 回复“确认提交”。",
    "Command(resume) 批准中断，恢复线程并完成 Mock 提交。",
], "Deep Agent 的重点不是一次生成结果，而是线程能在危险操作前暂停。用户确认后，不需要重新执行整个流程，而是从 checkpoint 恢复。", demo=[
    "进入 ai_deep_agents_assistant 目录。",
    "执行 .\\start_windows.ps1 -Port 8020。",
    "打开 http://127.0.0.1:8020/docs。",
    "使用 session_id：demo-deep-001。",
    "发送完整报销描述。",
    "展示预览、interrupt 和 awaiting_confirmation。",
    "继续使用 demo-deep-001 发送：确认提交。",
], note="最终提交进入 Mock 审批服务，不是生产审批平台。")

add_slide(doc, 14, "预期效果与当前成果", "当前完成的是功能验证和架构验证", [
    "用户能够通过自然语言表达审批或日报目标。",
    "主项目当前已经实现请假审批和日报场景。",
    "日报已经覆盖内容收集、上下文加载、草稿保存、预览和人工确认。",
    "不同实现可以返回 Agent 身份、状态、trace、preview 或 messages。",
    "Deep Agent 项目验证了 checkpoint 和人工确认机制。",
    "下一步需要量化平均完成轮次、工具调用成功率、响应时间、Token 成本和失败重试率。",
], "目前还没有真实的效率统计，所以不能说效率提升百分之多少。当前成果是证明问答式体验可行，并验证不同架构的安全边界和可观测性。")

add_slide(doc, 15, "遇到的问题与最终结论", "真正困难的是状态、边界和稳定性", [
    "Flow 节点膨胀：规则越多，条件边和修改分支越复杂。",
    "Agent 不确定性：工具顺序、循环次数、Token 和延迟难预测。",
    "意图区分困难：取消提交、修改内容和取消整个流程不能混淆。",
    "会话状态串线：复用 session 可能继续进入上一种 Agent。",
    "演示要求：三种日报实现必须使用不同 session；同一流程的确认必须使用原 session。",
    "Demo 与生产边界：规则模拟、Demo 保存和 Mock 后端必须明确说明。",
    "最终结论：AI 负责理解，代码负责边界。",
], "最后回到最初的问题：用户希望用一句话完成任务，但业务系统不能因为使用 AI 就失去规则和安全边界。最终要做的不是给 AI 最大的决策权，而是给它合适的决策权。")

doc.core_properties.title = "AI效能分享-审批与日报四种AI实现-逐页内容稿"
doc.core_properties.subject = "PPT每页页面标题、页面内容和讲解稿"
doc.core_properties.author = ""
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)

with ZipFile(OUTPUT) as archive:
    error = archive.testzip()
if error:
    raise RuntimeError(error)

check = Document(OUTPUT)
slide_headings = [p.text for p in check.paragraphs if p.text.startswith("第 ") and p.text.endswith(" 页")]
print({"output": str(OUTPUT), "size": OUTPUT.stat().st_size, "slides": len(slide_headings), "paragraphs": len(check.paragraphs)})
