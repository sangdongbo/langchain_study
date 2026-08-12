from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\PythonProject\LearnOne\AI效能分享-审批与日报四种AI实现-PPT生成文档.docx")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
CYAN = RGBColor(0, 151, 178)
TEAL = RGBColor(0, 121, 107)
GOLD = RGBColor(181, 129, 0)
RED = RGBColor(155, 28, 28)
GRAY = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 100), ("bottom", 100), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def set_run(run, size=11, color=None, bold=False, italic=False, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", size=11, color=None, bold=False, italic=False, after=6, before=0, alignment=None, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.2
    if alignment is not None:
        paragraph.alignment = alignment
    if text:
        run = paragraph.add_run(text)
        set_run(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(text), size=10.5)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(text), size=10.5)
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label}：")
    set_run(label_run, size=10.5, color=label_color, bold=True)
    set_run(paragraph.add_run(text), size=10.5)
    add_paragraph(doc, after=3)


def add_h1(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run(run, size=17, color=BLUE, bold=True)
    return paragraph


def add_h2(doc, text):
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run(run, size=13, color=NAVY, bold=True)
    return paragraph


def add_slide_section(doc, number, title, conclusion, visible, visual, talk, demo=None, boundary=None):
    add_h1(doc, f"第 {number} 页｜{title}")
    add_callout(doc, "本页结论", conclusion, LIGHT_BLUE, BLUE)
    add_h2(doc, "屏幕文案")
    for item in visible:
        add_bullet(doc, item)
    add_h2(doc, "视觉建议")
    add_paragraph(doc, visual, size=10.5, color=GRAY)
    if demo:
        add_h2(doc, "现场演示动作")
        for item in demo:
            add_number(doc, item)
    if boundary:
        add_callout(doc, "真实性边界", boundary, LIGHT_GOLD, GOLD)
    add_h2(doc, "讲解要点")
    add_paragraph(doc, talk, size=10.5)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("AI 效能分享｜PPT 生成内容文档"), size=9, color=GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

add_paragraph(doc, "AI 效能分享会", size=12, color=CYAN, bold=True, before=24, after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "从复杂表单到一句话提交", size=28, color=NAVY, bold=True, after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "审批助手与日报的四种 AI 实现", size=16, color=BLUE, after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Flow · Agent Flow · Agentic Workflow · Deep Agent", size=11, color=GRAY, after=44, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_callout(doc, "用途", "本文件可直接上传到 AI PPT 生成网站，也可以复制开头的生成提示词，再将逐页内容作为素材粘贴。建议生成 15 页、16:9、约 20 分钟的中文演示文稿。", LIGHT_BLUE, BLUE)
add_paragraph(doc, "内容基于本地项目实际代码整理，已区分真实业务能力、规则模拟、Demo 状态和 Mock 后端，避免在分享中夸大效果。", size=10, color=GRAY, italic=True, after=24)

add_h1(doc, "一、可直接粘贴给 AI PPT 网站的生成提示词")
prompt = (
    "请生成一份 15 页、16:9、约 20 分钟的中文企业内部 AI 应用分享 PPT。主题为《从复杂表单到一句话提交——审批助手与日报的四种 AI 实现》。"
    "受众是公司内部业务和技术同事，表达要通俗、可信、适合现场讲解，避免写成技术说明书。整套使用统一的现代商务科技模板，建议深蓝色背景，青色、绿色、琥珀色、紫色分别代表固定 Flow、自主 Agent、Agentic Workflow、Deep Agent。"
    "每页只保留一句核心结论和少量关键词，优先使用时间线、流程图、三段式结构和横向对比，避免密集表格、卡片堆叠和大段代码。四张演示页必须突出‘用户输入什么、现场看什么、最后说明什么’。"
    "内容必须覆盖应用场景、实施路径、预期效果、遇到的问题。应用场景是解决平台审批入口多、表单字段复杂的问题，让用户通过自然语言问答完成审批或日报提交。"
    "实施路径依次为固定 Flow、Agent Flow、自主 Agent 与 Workflow 结合、Deep Agent。三个日报实现使用同一目标对比；Deep Agent 是独立审批场景，不做完全公平的性能比较。"
    "必须真实说明：Agentic Workflow 当前规划和正文生成是规则模拟，保存和提交是 Demo 状态；Deep Agent 使用真实 checkpoint 和人工中断机制，但业务后端是 Mock 服务。"
    "最后结论为：AI 负责理解，代码负责边界；让 AI 获得合适的决策权，而不是最大的决策权。"
)
add_callout(doc, "生成提示词", prompt, LIGHT_GRAY, NAVY)

add_h1(doc, "二、分享的核心故事")
add_paragraph(doc, "这次分享不是四个框架的功能介绍，而是一个业务需求经过四次实现后，对“执行权”不断重新分配的过程。", size=12, color=NAVY, bold=True, after=10)
for item in [
    "业务目标始终不变：用户用一句自然语言表达审批或日报需求。",
    "固定 Flow 解决流程能否稳定执行的问题。",
    "自主 Agent 验证模型能否自己选择工具、完成目标。",
    "Agentic Workflow 解决 AI 如何安全进入真实业务流程的问题。",
    "Deep Agent 探索更复杂、更长线程任务中的 checkpoint 和人工确认。",
]:
    add_bullet(doc, item)
add_callout(doc, "主线", "不是框架越新越好，而是业务风险越高，代码应该保留越多的执行控制权。", LIGHT_GOLD, GOLD)

add_slide_section(doc, 1, "封面：从复杂表单到一句话提交", "这次实践的目标，是让用户不再先学习表单和系统入口。", [
    "主标题：从复杂表单到一句话提交",
    "副标题：审批助手与日报的四种 AI 实现",
    "底部关键词：Flow / Agent Flow / Agentic Workflow / Deep Agent",
], "使用简洁封面。左侧大标题，右侧可以放抽象对话流、表单或节点网络视觉；不要放目录和大段说明。", "开场不要先介绍 LangGraph、Agent 或 Deep Agent。先说：平台里的审批入口和字段很多，我希望用户只需要说一句话，系统负责理解、追问、生成预览，用户只负责最终确认。")

add_slide_section(doc, 2, "分享路线", "先讲为什么做，再用四次现场演示说明实现路径。", [
    "应用场景：为什么要做问答式审批助手",
    "实施路径：从 Flow 到 Agent，再到混合架构",
    "三种日报：同一目标的不同执行方式",
    "Deep Agent：审批场景中的长线程与人工确认",
    "预期效果与遇到的问题",
], "使用一条纵向或横向路线，不要使用传统六行目录表。", "说明内容正好覆盖通知要求的应用场景、实施路径、预期效果和遇到的问题，四次演示穿插在主线中间。")

add_slide_section(doc, 3, "应用场景：为什么需要问答式审批", "用户不应该先学会表单，才能完成审批。", [
    "现在：找入口 → 选模板 → 理解字段 → 逐项填写 → 找审批人 → 提交",
    "希望：直接说‘我要请假三天’或‘把今天完成的工作写进日报’",
    "AI 负责理解和追问；业务代码负责校验、预览和写入",
], "做明显的 Before / After。左侧表现复杂表单路径，右侧突出一句自然语言和生成预览。", "强调这是为了解决用户记忆入口、理解字段和重复填写的成本，不是为了展示模型能力。")

add_slide_section(doc, 4, "实施路径：执行权的演进", "四次演进的核心，是不断调整模型和代码分别做什么决定。", [
    "固定 Flow：代码决定节点、条件边和执行顺序",
    "自主 Agent：模型决定调用什么工具和下一步",
    "Agentic Workflow：Agent 负责理解生成，Workflow 负责业务 Gate",
    "Deep Agent：目标驱动、checkpoint、人工确认和线程恢复",
], "使用一条四阶段时间线，每种实现固定一种颜色。", "固定 Flow 解决稳定性，自主 Agent 验证自主工具调用，Agentic Workflow 尝试安全进入业务，Deep Agent 探索更复杂任务。")

add_slide_section(doc, 5, "固定 Flow：daily_report_agent", "流程顺序完全由 Workflow 控制，适合规则明确、写入风险高的业务。", [
    "入口 → 动作识别 → 日期 / 内容收集 → 加载上下文",
    "保存草稿 → 生成预览 → 等待确认",
    "确认提交 / 修改日期 / 修改内容 / 取消",
    "优势：路径确定、状态清晰、异常容易定位和测试",
], "使用从左到右的固定流程图，节点不超过 6 个，可将日期和内容收集合并展示。", "不要逐个念节点名。用业务语言说明系统先理解动作、补齐信息、保存草稿、展示预览，最后等待确认。")

add_slide_section(doc, 6, "现场演示 1：固定 Flow", "最终结果之外，重点看路径是否稳定、失败是否能定位。", [
    "输入：写日报：今天完成审批助手开发",
    "观察：是否进入 daily_report_agent",
    "观察：trace 是否显示固定节点顺序",
    "观察：是否先生成预览，再等待确认",
], "页面设计为三段：左侧输入，中间 trace，右侧预览和确认。", "展示固定路径、预览和人工确认。失败时也能从 trace 直接知道卡在哪一步。", demo=[
    "调用 8010 的 POST /api/ai-approval/chat。",
    "使用 session_id：demo-flow-001。",
    "发送‘写日报：今天完成审批助手开发’。",
    "展示 agent、status、preview 和 trace。",
    "继续使用相同 session_id 回复‘确认提交’。",
])

add_slide_section(doc, 7, "自主 Agent：daily_report_create_agent", "模型获得工具选择权，但提交仍必须保留安全守卫。", [
    "create_agent / ReAct Agent 自主选择下一步",
    "可调用日期、上下文、保存草稿、预览等日报工具",
    "提交工具要求 confirmed=true，未确认不能提交",
    "状态主要保存在 daily_report_agent_messages",
    "代价：调用顺序、循环次数、Token 和延迟更难预测",
], "表现为用户目标进入 Agent，Agent 在工具列表中自主选择，再观察工具结果。", "自主 Agent 的问题不是最后结果能不能做出来，而是过程由模型决定后，状态排查和成本控制变得更难。")

add_slide_section(doc, 8, "现场演示 2：自主 Agent", "同样生成日报，但排查方式已经从节点状态变成消息和工具调用历史。", [
    "输入：自主agent写日报：今天完成审批助手开发2",
    "观察：daily_report_agent_messages",
    "观察：模型选择了哪些工具、工具调用顺序是什么",
    "结论：更灵活，但更难预测和复现",
], "使用‘输入 → messages → tool calls → 结果’布局，工具调用可以用简化标签表现。", "不要只展示最终回答，要主动展示 Agent 消息和 tool calls，说明业务阶段需要从历史中反推。", demo=[
    "使用新的 session_id：demo-agent-001。",
    "发送‘自主agent写日报：今天完成审批助手开发2’。",
    "查看 AgentMessage、ToolMessage 和工具参数。",
    "观察日期、上下文、草稿和预览工具的实际顺序。",
], boundary="该实现需要真实模型或 API。模型不可用时，不要包装成固定流程已经成功执行。")

add_slide_section(doc, 9, "Agentic Workflow：daily_report_agentic_workflow_demo", "Agent 负责理解和生成，Workflow 负责业务顺序和写入权限。", [
    "Agent：提取日期、工作内容，生成计划和日报正文",
    "Workflow：确认日期、加载上下文、保存草稿、生成预览",
    "Gate：没有预览不能提交；没有明确确认不能提交",
    "结构化状态：plan / context / compose / events / draft_saved",
], "使用三段式：Agent 理解 → Workflow 执行 → Gate 裁决。", "这是当前最值得继续完善的方向：既保留 AI 的理解能力，又让代码掌握真实写入顺序和权限。", boundary="当前 demo_agent_plan 和 demo_agent_compose 是规则模拟；加载上下文调用现有日报工具；保存草稿和最终提交使用 Demo 状态。")

add_slide_section(doc, 10, "现场演示 3：Agentic Workflow", "重点看 Agent 和 Workflow 的职责是否分开，以及 Gate 是否真正拦截提交。", [
    "输入：agentic workflow写日报：今天完成审批助手开发3",
    "观察：plan / context / compose 是否分开保存",
    "观察：trace 是否进入 demo_preview_gate",
    "确认后：是否进入 demo_submit_gate",
], "页面设计为‘结构化状态、Gate、结论’三栏，避免展示完整 JSON。", "展示 trace 和结构化状态。先说明未确认时不能提交，再用同一 session 确认。", demo=[
    "使用新的 session_id：demo-agentic-001。",
    "发送‘agentic workflow写日报：今天完成审批助手开发3’。",
    "展示 demo_agent_plan → demo_confirm_date → demo_load_context → demo_agent_compose → demo_save_draft → demo_preview_gate。",
    "继续使用相同 session_id 回复‘确认提交’。",
], boundary="最终 Demo 编号和提交结果用于验证架构，不代表已经真实写入生产 ERP。")

add_slide_section(doc, 11, "三种日报实现对比", "同一日报目标下，最大的区别是执行顺序由谁控制。", [
    "固定 Flow：Workflow 控制；状态最清晰；成本更可控；适合稳定主流程",
    "自主 Agent：Agent 控制；需要从 messages 和 tool calls 反推；适合工具实验",
    "Agentic Workflow：Agent 建议、Workflow 决定；状态分层；适合受控智能业务",
    "选择原则：写入风险越高，代码控制越多；任务越开放，Agent 权限可以越大",
], "不要使用密集 Excel 表格。建议画三条横向赛道，比较执行权、可观测性、成本和定位。", "不是越智能越好。生产业务首先要保证可定位、可恢复、可控制，然后再增加模型的决策空间。")

add_slide_section(doc, 12, "Deep Agent：ai_deep_agents_assistant", "Deep Agent 用于探索更复杂目标、长线程状态和 Human-in-the-loop。", [
    "create_deep_agent + MemorySaver checkpoint",
    "工具：用户上下文、审批模板、草稿、预览、提交",
    "submit_approval_request 前配置 interrupt_on",
    "确认后通过 Command(resume) 恢复线程",
    "支持请假、报销、采购等 Mock 审批场景",
], "表现为目标进入 Deep Agent，经过多个工具后在危险提交前暂停，再由人工批准恢复。", "Deep Agent 与前三个不是同一日报场景，所以不做完全公平的性能排名。它的价值是展示长线程、checkpoint 和人工中断。", boundary="Deep Agents、checkpoint 和 interrupt_on 是真实框架能力；审批业务后端是确定性的 Mock 服务。")

add_slide_section(doc, 13, "现场演示 4：Deep Agent", "重点不是一次回答，而是线程能够在危险工具前暂停并从 checkpoint 恢复。", [
    "输入：我要报销差旅费，金额 5200 元，因为去上海拜访客户，发票已提供",
    "第一轮：Agent 调用用户、模板和草稿工具，生成审批预览",
    "中断：submit_approval_request 前触发 interrupt_on",
    "确认：同一 session 回复‘确认提交’，恢复并完成 Mock 提交",
], "三段式：第一轮工具调用 → 人工中断 → 确认后恢复。", "现场要明确指出 awaiting_confirmation、interrupt 和恢复后的 submitted 状态。", demo=[
    "启动 ai_deep_agents_assistant：.\\start_windows.ps1 -Port 8020。",
    "打开 8020 的 Swagger。",
    "使用 session_id：demo-deep-001。",
    "发送完整报销描述。",
    "展示 interrupt 和预览。",
    "继续使用相同 session_id 回复‘确认提交’。",
], boundary="最终提交写入 Mock 审批服务，不是生产审批平台。")

add_slide_section(doc, 14, "预期效果与当前成果", "当前成果是功能和架构验证，不能虚构尚未测量的效率数字。", [
    "用户可以用自然语言表达审批或日报目标",
    "主项目目前实现请假和日报能力",
    "日报已经覆盖内容收集、草稿、预览和人工确认",
    "响应中可以观察 Agent 身份、状态、trace、preview 或 messages",
    "Deep Agent 项目验证 checkpoint 和人工确认机制",
    "下一步指标：平均完成轮次、工具成功率、响应时间、Token、失败重试率",
], "左侧列出已经实现，右侧列出下一步需要量化的指标。", "不要说效率提升 80% 等没有测量依据的数字。当前可以证明的是体验、架构和安全边界。")

add_slide_section(doc, 15, "遇到的问题与结论", "真正困难的不是调用模型，而是状态、边界和演示稳定性。", [
    "Flow 节点膨胀：规则、条件边和修改分支越来越重",
    "Agent 不确定性：工具顺序、循环次数、Token 和延迟难预测",
    "意图语义：取消提交、修改内容、取消整个流程必须区分",
    "会话串线：复用 session 可能继续路由到上一种 Agent",
    "演示稳定性：三个日报实验必须使用不同 session；同一流程确认必须复用原 session",
    "Demo 与生产边界：规则模拟、Demo 保存和 Mock 后端必须如实说明",
    "最终结论：AI 负责理解，代码负责边界",
], "建议使用问题清单收束，并用一句大字结论结束，不要再放复杂架构图。", "最后回到开场：用户只想用一句话完成任务，但后台不能因此失去业务规则。让 AI 获得合适的决策权，而不是最大的决策权。")

add_h1(doc, "三、四种实现的技术对比摘要")
table = doc.add_table(rows=1, cols=4)
set_table_geometry(table, [1900, 2300, 2660, 2500])
headers = ["实现", "执行权", "状态与排查", "适用定位"]
for index, value in enumerate(headers):
    cell = table.rows[0].cells[index]
    set_cell_shading(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(paragraph.add_run(value), size=10, color=NAVY, bold=True)
rows = [
    ("daily_report_agent", "Workflow 决定顺序", "节点状态与 trace 清晰", "强规则生产流程"),
    ("daily_report_create_agent", "Agent 自主选工具", "从 messages / tool calls 反推", "自主工具调用实验"),
    ("daily_report_agentic_workflow_demo", "Agent 建议，Workflow Gate", "计划、上下文、生成、Gate 分层", "受控智能流程"),
    ("ai_deep_agents_assistant", "目标驱动 + HITL", "Checkpoint、interrupt、resume", "长线程和复杂目标探索"),
]
for row_data in rows:
    row = table.add_row()
    for index, value in enumerate(row_data):
        paragraph = row.cells[index].paragraphs[0]
        set_run(paragraph.add_run(value), size=9.5)
set_table_geometry(table, [1900, 2300, 2660, 2500])

add_h1(doc, "四、现场演示准备清单")
for item in [
    "提前启动 8010 日报/审批服务，并确认 Swagger 可以访问。",
    "提前启动 8020 Deep Agent 服务，并确认模型配置可用。",
    "为三个日报实验分别准备 demo-flow-001、demo-agent-001、demo-agentic-001。",
    "Deep Agent 使用 demo-deep-001。",
    "同一流程的确认提交必须继续使用原 session_id。",
    "准备四组成功截图或短录屏，作为模型或网络异常时的兜底。",
    "演示时优先展示 input、agent、status、trace/messages、preview，不展示完整大段 JSON。",
    "明确讲出 Agentic Workflow 的 Demo 状态和 Deep Agent 的 Mock 后端。",
]:
    add_bullet(doc, item)

add_h2(doc, "建议现场提问文本")
for text_value in [
    "固定 Flow：写日报：今天完成审批助手开发",
    "自主 Agent：自主agent写日报：今天完成审批助手开发2",
    "Agentic Workflow：agentic workflow写日报：今天完成审批助手开发3",
    "Deep Agent：我要报销差旅费，金额 5200 元，因为去上海拜访客户，发票已提供",
    "确认动作：确认提交",
]:
    add_bullet(doc, text_value)

add_h1(doc, "五、技术依据与真实性边界")
for item in [
    r"固定 Flow：D:\PythonProject\LearnOne\ai_approval_assistant\app\graph\daily_report_workflow.py",
    r"自主 Agent：D:\PythonProject\LearnOne\ai_approval_assistant\app\agents\daily_report_create_agent.py",
    r"日报工具：D:\PythonProject\LearnOne\ai_approval_assistant\app\tools\daily_report_tools.py",
    r"Agentic Workflow：D:\PythonProject\LearnOne\ai_approval_assistant\app\graph\daily_report_agentic_workflow_demo.py",
    r"Agentic 节点：D:\PythonProject\LearnOne\ai_approval_assistant\app\agents\daily_report_agentic_workflow_demo.py",
    r"Deep Agent：D:\PythonProject\LearnOne\ai_deep_agents_assistant\app\agents\approval_agent.py",
    r"Deep Agent 工具与服务：D:\PythonProject\LearnOne\ai_deep_agents_assistant\app\tools\approval_tools.py；app\services\approval_service.py",
]:
    add_bullet(doc, item)
add_callout(doc, "最终表达原则", "把实现边界讲清楚比包装成‘已经全面落地’更有说服力。分享重点是实践过程、架构取舍和真实问题。", LIGHT_GOLD, GOLD)

doc.core_properties.title = "AI效能分享-审批与日报四种AI实现-PPT生成文档"
doc.core_properties.subject = "AI应用分享会PPT生成内容与现场演示脚本"
doc.core_properties.author = ""

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
